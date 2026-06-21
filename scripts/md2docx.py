#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
md2docx.py — 将 Markdown 论文初稿转换为 Word 文档
支持：标题、表格、列表、粗体/斜体、代码块、引用、附录
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 中文字体设置
def set_cn_font(run, font_name="宋体", size=10.5, bold=False, italic=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def add_para_with_runs(doc, segments, size=10.5, bold=False, italic=False, first_line_indent=True):
    """添加段落，可包含多种格式文本段。segments: [(text, {bold,italic}), ...]"""
    p = doc.add_paragraph()
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    for text, fmt in segments:
        run = p.add_run(text)
        is_bold = fmt.get("bold", bold)
        is_italic = fmt.get("italic", italic)
        set_cn_font(run, size=size, bold=is_bold, italic=is_italic)
    return p


def parse_inline(text):
    """解析行内格式：[('**bold**', 'bold'), ('normal', 'normal')] -> [(text, fmt_dict), ...]"""
    segments = []
    # 处理 **bold** 和 *italic* 和 `code`
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            segments.append((text[pos:m.start()], {}))
        token = m.group(0)
        if token.startswith("**"):
            segments.append((token[2:-2], {"bold": True}))
        elif token.startswith("`"):
            segments.append((token[1:-1], {"code": True}))
        elif token.startswith("*"):
            segments.append((token[1:-1], {"italic": True}))
        pos = m.end()
    if pos < len(text):
        segments.append((text[pos:], {}))
    if not segments:
        segments = [(text, {})]
    return segments


def add_inline_para(doc, text, size=10.5, bold=False, italic=False, first_line_indent=True, justify=True):
    """添加带行内格式的段落"""
    p = doc.add_paragraph()
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    segments = parse_inline(text)
    for seg_text, fmt in segments:
        run = p.add_run(seg_text)
        is_bold = fmt.get("bold", bold)
        is_italic = fmt.get("italic", italic)
        is_code = fmt.get("code", False)
        if is_code:
            set_cn_font(run, font_name="Consolas", size=size, bold=is_bold)
        else:
            set_cn_font(run, size=size, bold=is_bold, italic=is_italic)
    return p


def add_heading(doc, text, level=1):
    """添加标题"""
    sizes = {1: 16, 2: 14, 3: 12, 4: 11}
    size = sizes.get(level, 11)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(text)
    set_cn_font(run, font_name="黑体" if level <= 2 else "宋体",
                size=size, bold=True)
    # 标题加 outline level
    pPr = p._p.get_or_add_pPr()
    outlineLvl = OxmlElement("w:outlineLvl")
    outlineLvl.set(qn("w:val"), str(level - 1))
    pPr.append(outlineLvl)
    return p


def add_code_block(doc, code_text):
    """添加代码块（等宽字体 + 灰底）"""
    for line in code_text.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line if line else " ")
        set_cn_font(run, font_name="Consolas", size=9, italic=False)
    # 代码块后空一行
    doc.add_paragraph()


def add_quote(doc, text):
    """添加引用块"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_cn_font(run, size=10, italic=True)
    # 加灰底
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p


def add_list_item(doc, text, ordered=False, level=0):
    """添加列表项"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(0.74 + level * 0.6)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    p.paragraph_format.space_after = Pt(0)
    segments = parse_inline(text)
    for seg_text, fmt in segments:
        run = p.add_run(seg_text)
        is_bold = fmt.get("bold", False)
        is_italic = fmt.get("italic", False)
        is_code = fmt.get("code", False)
        if is_code:
            set_cn_font(run, font_name="Consolas", size=10)
        else:
            set_cn_font(run, size=10.5, bold=is_bold, italic=is_italic)
    return p


def add_table(doc, header, rows):
    """添加表格"""
    n_cols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"

    # 表头
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_cn_font(run, font_name="黑体", size=10, bold=True)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 灰底
        tcPr = hdr_cells[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D9E2F3")
        tcPr.append(shd)

    # 数据行
    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row):
            row_cells[c_idx].text = ""
            p = row_cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.2
            # 解析行内格式
            segments = parse_inline(str(cell_text))
            for seg_text, fmt in segments:
                run = p.add_run(seg_text)
                is_bold = fmt.get("bold", False)
                is_italic = fmt.get("italic", False)
                is_code = fmt.get("code", False)
                if is_code:
                    set_cn_font(run, font_name="Consolas", size=9)
                else:
                    set_cn_font(run, size=9.5, bold=is_bold, italic=is_italic)
            row_cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph()  # 表格后空行


def add_figure_placeholder(doc, caption):
    """添加图片占位框"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("[ 此处插入图片 ]")
    set_cn_font(run, size=10, italic=True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(8)
    run2 = p2.add_run(caption)
    set_cn_font(run2, size=9.5, italic=True, bold=True)
    return p


def add_horizontal_rule(doc):
    """添加分隔线"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)


def parse_table_block(lines, start_idx):
    """解析 markdown 表格块，返回 (header, rows, 结束行号)"""
    header = [c.strip() for c in lines[start_idx].strip().strip("|").split("|")]
    rows = []
    i = start_idx + 2  # 跳过分隔行
    while i < len(lines) and lines[i].strip().startswith("|"):
        row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        rows.append(row)
        i += 1
    return header, rows, i


def convert_md_to_docx(md_path, docx_path):
    md_text = Path(md_path).read_text(encoding="utf-8")
    lines = md_text.split("\n")

    doc = Document()

    # 设置默认样式
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    i = 0
    in_code_block = False
    code_buffer = []
    in_table = False
    table_buffer = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 代码块
        if stripped.startswith("```"):
            if in_code_block:
                add_code_block(doc, "\n".join(code_buffer))
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # 标题
        if stripped.startswith("# "):
            add_heading(doc, stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            add_heading(doc, stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            add_heading(doc, stripped[4:].strip(), level=3)
        elif stripped.startswith("#### "):
            add_heading(doc, stripped[5:].strip(), level=4)

        # 水平线
        elif stripped == "---":
            add_horizontal_rule(doc)

        # 引用
        elif stripped.startswith("> "):
            add_quote(doc, stripped[2:].strip())
        elif stripped == ">":
            add_quote(doc, "")

        # 表格
        elif stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s\-:|]+\|$", lines[i + 1].strip()):
            # 收集整个表格
            table_lines = []
            j = i
            while j < len(lines) and lines[j].strip().startswith("|"):
                table_lines.append(lines[j])
                j += 1
            if len(table_lines) >= 2:
                header = [c.strip() for c in table_lines[0].strip().strip("|").split("|")]
                rows = []
                for tl in table_lines[2:]:
                    row = [c.strip() for c in tl.strip().strip("|").split("|")]
                    rows.append(row)
                add_table(doc, header, rows)
            i = j
            continue

        # 列表
        elif re.match(r"^\s*-\s+", line):
            text = re.sub(r"^\s*-\s+", "", line)
            add_list_item(doc, text, ordered=False)
        elif re.match(r"^\s*\d+\.\s+", line):
            text = re.sub(r"^\s*\d+\.\s+", "", line)
            add_list_item(doc, text, ordered=True)

        # 普通段落
        elif stripped and not stripped.startswith("|") and not stripped.startswith(">"):
            # 合并连续非空行为一段
            buf = [stripped]
            j = i + 1
            while j < len(lines) and lines[j].strip() and not lines[j].strip().startswith(("#", "|", ">", "```", "-", "---")) and not re.match(r"^\s*\d+\.\s+", lines[j]):
                buf.append(lines[j].strip())
                j += 1
            para_text = " ".join(buf)
            # 判断是否是图标题（图 N xxx）
            is_fig_caption = para_text.startswith("图 ") or "**图" in para_text[:6]
            if is_fig_caption:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(8)
                run = p.add_run(para_text)
                set_cn_font(run, size=9.5, bold=True)
            else:
                add_inline_para(doc, para_text, size=10.5)
            i = j
            continue

        i += 1

    doc.save(docx_path)
    print(f"[OK] 已生成: {docx_path}")


if __name__ == "__main__":
    md_file = r"c:\Users\Lenovo\Desktop\灵境制造（上线版）\docs\research\pi-lnn-mssp-draft-v0.1.md"
    docx_file = r"c:\Users\Lenovo\Desktop\灵境制造（上线版）\docs\research\pi-lnn-mssp-draft-v0.1.docx"
    convert_md_to_docx(md_file, docx_file)
