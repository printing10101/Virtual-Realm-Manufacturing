"""PDF/Word 文档文本提取模块（M1.4）。

支持从 PDF 和 Word 格式文档中提取文本内容，保持页面级别的逻辑结构。
对于扫描版 PDF，提供 OCR 降级处理。

依赖：
    - pdfplumber（PDF 文本提取，主路径）
    - pdfminer.six（PDF 文本提取，备用路径）
    - python-docx（Word 文档提取）
    - pytesseract + pdf2image（OCR 降级，可选）
"""

from __future__ import annotations

import json
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# 数据类
# ---------------------------------------------------------------------------


@dataclass
class DocumentPage:
    """单页文档内容。"""

    page_number: int
    text: str
    source: str = ""


@dataclass
class ExtractedDocument:
    """提取后的文档结构。"""

    source_path: str
    pages: list[DocumentPage] = field(default_factory=list)
    total_pages: int = 0
    extraction_method: str = ""
    metadata: dict[str, Any] = field(default_factory=dict)

    @property
    def full_text(self) -> str:
        """获取全部文本内容。"""
        return "\n\n".join(p.text for p in self.pages if p.text.strip())


# ---------------------------------------------------------------------------
# PDF 提取器
# ---------------------------------------------------------------------------


class PDFExtractor:
    """PDF 文档文本提取器。

    提取策略（按优先级）：
        1. pdfplumber 文本层提取（高质量）
        2. pdfminer.six 文本层提取（备用）
        3. OCR 识别（扫描版 PDF 降级方案，需安装 pytesseract）
    """

    def extract(self, file_path: str | Path) -> ExtractedDocument:
        """提取 PDF 文档内容。

        Args:
            file_path: PDF 文件路径。

        Returns:
            ExtractedDocument 包含按页组织的文本内容。

        Raises:
            FileNotFoundError: 文件不存在。
            ValueError: 文件为空或提取失败。
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"PDF 文件不存在: {file_path}")

        pages = self._extract_with_pdfplumber(file_path)
        method = "pdfplumber"

        if not pages or all(not p.text.strip() for p in pages):
            logger.info("pdfplumber 提取结果为空，尝试 pdfminer")
            pages = self._extract_with_pdfminer(file_path)
            method = "pdfminer"

        if not pages or all(not p.text.strip() for p in pages):
            logger.info("文本层提取均为空，尝试 OCR")
            pages = self._extract_with_ocr(file_path)
            method = "ocr"

        if not pages:
            raise ValueError(f"无法从 PDF 提取文本: {file_path}")

        return ExtractedDocument(
            source_path=str(file_path),
            pages=pages,
            total_pages=len(pages),
            extraction_method=method,
            metadata={"filename": file_path.name},
        )

    @staticmethod
    def _extract_with_pdfplumber(file_path: Path) -> list[DocumentPage]:
        """使用 pdfplumber 提取文本。"""
        try:
            import pdfplumber
        except ImportError:
            logger.warning("pdfplumber 未安装，跳过")
            return []

        pages: list[DocumentPage] = []
        try:
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text() or ""
                    pages.append(
                        DocumentPage(
                            page_number=i + 1,
                            text=text.strip(),
                            source=str(file_path),
                        )
                    )
        except (OSError, ValueError, TypeError) as exc:
            logger.warning("pdfplumber 提取失败: %s", exc)
        return pages

    @staticmethod
    def _extract_with_pdfminer(file_path: Path) -> list[DocumentPage]:
        """使用 pdfminer.six 提取文本。"""
        try:
            from pdfminer.high_level import extract_text_to_fp
            from pdfminer.layout import LAParams
            from io import StringIO
        except ImportError:
            logger.warning("pdfminer.six 未安装，跳过")
            return []

        pages: list[DocumentPage] = []
        try:
            laparams = LAParams(
                line_margin=0.5,
                word_margin=0.1,
                char_margin=2.0,
                boxes_flow=0.5,
            )
            output = StringIO()
            with open(file_path, "rb") as f:
                extract_text_to_fp(f, output, laparams=laparams, codec="utf-8")
            raw_text = output.getvalue()
            # 按分页符分割
            raw_pages = raw_text.split("\x0c")
            for i, page_text in enumerate(raw_pages):
                text = page_text.strip()
                if text:
                    pages.append(
                        DocumentPage(
                            page_number=i + 1,
                            text=text,
                            source=str(file_path),
                        )
                    )
        except (OSError, ValueError, TypeError) as exc:
            logger.warning("pdfminer 提取失败: %s", exc)
        return pages

    @staticmethod
    def _extract_with_ocr(file_path: Path) -> list[DocumentPage]:
        """使用 OCR 提取扫描版 PDF（需 pytesseract + pdf2image）。"""
        try:
            import pytesseract
            from pdf2image import convert_from_path
        except ImportError:
            logger.warning("OCR 依赖未安装（需要 pytesseract + pdf2image），跳过 OCR")
            return []

        pages: list[DocumentPage] = []
        try:
            images = convert_from_path(str(file_path))
            for i, image in enumerate(images):
                text = pytesseract.image_to_string(image, lang="chi_sim+eng")
                pages.append(
                    DocumentPage(
                        page_number=i + 1,
                        text=text.strip(),
                        source=str(file_path),
                    )
                )
        except (OSError, ValueError, TypeError) as exc:
            logger.warning("OCR 提取失败: %s", exc)
        return pages


# ---------------------------------------------------------------------------
# Word 提取器
# ---------------------------------------------------------------------------


class WordExtractor:
    """Word 文档（.docx）文本提取器。"""

    def extract(self, file_path: str | Path) -> ExtractedDocument:
        """提取 Word 文档内容。

        Args:
            file_path: Word 文件路径（.docx）。

        Returns:
            ExtractedDocument 包含文档内容。

        Raises:
            FileNotFoundError: 文件不存在。
            ValueError: 提取失败。
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"Word 文件不存在: {file_path}")

        pages = self._extract_with_docx(file_path)
        method = "python-docx"

        if not pages or all(not p.text.strip() for p in pages):
            pages = self._extract_with_docx2txt(file_path)
            method = "docx2txt"

        if not pages:
            raise ValueError(f"无法从 Word 文档提取文本: {file_path}")

        return ExtractedDocument(
            source_path=str(file_path),
            pages=pages,
            total_pages=len(pages),
            extraction_method=method,
            metadata={"filename": file_path.name},
        )

    @staticmethod
    def _extract_with_docx(file_path: Path) -> list[DocumentPage]:
        """使用 python-docx 提取文本。"""
        try:
            from docx import Document
        except ImportError:
            logger.warning("python-docx 未安装，跳过")
            return []

        try:
            doc = Document(str(file_path))
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            full_text = "\n\n".join(paragraphs)
            if not full_text.strip():
                return []

            return [
                DocumentPage(
                    page_number=1,
                    text=full_text,
                    source=str(file_path),
                )
            ]
        except (OSError, ValueError, TypeError) as exc:
            logger.warning("python-docx 提取失败: %s", exc)
            return []

    @staticmethod
    def _extract_with_docx2txt(file_path: Path) -> list[DocumentPage]:
        """使用 docx2txt 提取文本（备用）。"""
        try:
            import docx2txt
        except ImportError:
            logger.warning("docx2txt 未安装，跳过")
            return []

        try:
            text = docx2txt.process(str(file_path))
            if not text or not text.strip():
                return []
            return [
                DocumentPage(
                    page_number=1,
                    text=text.strip(),
                    source=str(file_path),
                )
            ]
        except (OSError, ValueError, TypeError) as exc:
            logger.warning("docx2txt 提取失败: %s", exc)
            return []


# ---------------------------------------------------------------------------
# 工厂函数与工具
# ---------------------------------------------------------------------------


def extract_document(file_path: str | Path) -> ExtractedDocument:
    """根据文件扩展名自动选择提取器。

    Args:
        file_path: 文档文件路径（支持 .pdf / .docx / .doc）。

    Returns:
        ExtractedDocument 提取结果。

    Raises:
        ValueError: 不支持的文件格式。
    """
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return PDFExtractor().extract(file_path)
    elif suffix in (".docx", ".doc"):
        return WordExtractor().extract(file_path)
    else:
        raise ValueError(f"不支持的文件格式: {suffix}（仅支持 .pdf / .docx / .doc）")


def chunk_pages(pages: list[DocumentPage], max_pages: int = 10) -> list[list[DocumentPage]]:
    """将页面列表按 max_pages 分块。

    Args:
        pages: 页面列表。
        max_pages: 每块最大页数（默认 10）。

    Returns:
        分块后的页面列表。
    """
    if not pages:
        return []
    return [pages[i : i + max_pages] for i in range(0, len(pages), max_pages)]


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------


def main():  # pragma: no cover
    """命令行文档提取工具。"""
    import argparse

    parser = argparse.ArgumentParser(description="从 PDF/Word 文档提取文本")
    parser.add_argument("input", help="输入文件路径")
    parser.add_argument("--output", "-o", help="输出 JSON 文件路径")
    args = parser.parse_args()

    logging.basicConfig(level=logging.INFO)

    result = extract_document(args.input)
    data = {
        "source_path": result.source_path,
        "total_pages": result.total_pages,
        "extraction_method": result.extraction_method,
        "pages": [{"page_number": p.page_number, "text": p.text} for p in result.pages],
        "full_text": result.full_text,
    }

    if args.output:
        Path(args.output).parent.mkdir(parents=True, exist_ok=True)
        with open(args.output, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        logger.info("已保存到 %s", args.output)
    else:
        # CLI 输出到 stdout，使用 print 是合理的
        print(json.dumps(data, ensure_ascii=False, indent=2))


if __name__ == "__main__":  # pragma: no cover
    main()
