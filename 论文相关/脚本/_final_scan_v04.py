"""论文1 v0.4 最终一致性扫描：确认所有 v0.3 残留已清除。"""
import docx

DOC_PATH = r'论文相关\论文与实验报告\论文1_DL-LNN颤振预测主论文.docx'
doc = docx.Document(DOC_PATH)

# v0.3 残留关键词
stale_patterns = [
    '0.987',           # 旧 PCC 值
    '略逊于',           # 旧精度声明
    '非最优',           # 旧精度声明
    '14.8%',           # 旧改善百分比
    '4.6%',            # 旧改善百分比
    '8.89%',           # 旧 DL-LNN 落后 PINN 百分比
    '1.60%',           # 旧 DL-LNN 落后 PINN 百分比
    '0.3744',          # 旧 Synthetic MAE
    '1.2061',          # 旧 Industrial MAE
    '设计目标并非',      # 旧双重优势声明
    'MAE 尚非最优',     # 旧结论
    '20.000',          # 旧 GP 发散值
    '20.000†',
]

print('=' * 70)
print('论文1 v0.4 最终一致性扫描')
print('=' * 70)

total_residuals = 0

# 扫描段落
print('\n[段落扫描]')
for i, p in enumerate(doc.paragraphs):
    for pattern in stale_patterns:
        if pattern in p.text:
            print(f'  [{i}] 发现 "{pattern}": {p.text[:120]}')
            total_residuals += 1

# 扫描表格
print('\n[表格扫描]')
for t_idx, t in enumerate(doc.tables):
    for r_idx, row in enumerate(t.rows):
        for c_idx, cell in enumerate(row.cells):
            for pattern in stale_patterns:
                if pattern in cell.text:
                    print(f'  表格[{t_idx}] 行[{r_idx}] 列[{c_idx}] 发现 "{pattern}": {cell.text[:80]}')
                    total_residuals += 1

# 扫描 v0.4 正确值是否已植入
print('\n[v0.4 正确值植入验证]')
v04_patterns = ['0.3222', '0.9289', '0.9953', '0.9968', '0.9680', '2.6367', '0.5076', '0.9560']
for pattern in v04_patterns:
    found_in_para = False
    found_in_table = False
    for p in doc.paragraphs:
        if pattern in p.text:
            found_in_para = True
            break
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if pattern in cell.text:
                    found_in_table = True
                    break
            if found_in_table:
                break
        if found_in_table:
            break
    status = '✓' if (found_in_para or found_in_table) else '✗'
    loc = []
    if found_in_para:
        loc.append('段落')
    if found_in_table:
        loc.append('表格')
    print(f'  {status} {pattern}: {"/".join(loc) if loc else "未找到"}')

print('\n' + '=' * 70)
if total_residuals == 0:
    print('✅ 论文1 v0.4 一致性完整，所有 v0.3 残留已清除')
else:
    print(f'⚠️  发现 {total_residuals} 处 v0.3 残留，需进一步修复')
print('=' * 70)
