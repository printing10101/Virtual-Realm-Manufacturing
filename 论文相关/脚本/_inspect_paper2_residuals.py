"""检查论文2 中 0.987 残留的上下文。"""
import docx

DOC_PATH = r'论文相关\论文与实验报告\论文2_PCC_Loss通用化方法论.docx'
doc = docx.Document(DOC_PATH)

print('=' * 70)
print('论文2 中 0.987 残留的上下文')
print('=' * 70)

# [176] 段落
p176 = doc.paragraphs[176]
print(f'\n[176] 段落完整内容:')
print(f'  {p176.text}')

# 表格[3] 行[3]
t3 = doc.tables[3]
print(f'\n表格[3] ({len(t3.rows)}行×{len(t3.columns)}列):')
for r_idx, row in enumerate(t3.rows):
    cells = [cell.text.strip() for cell in row.cells]
    print(f'  行[{r_idx}]: {cells}')

# 表格[5] 行[1]
t5 = doc.tables[5]
print(f'\n表格[5] ({len(t5.rows)}行×{len(t5.columns)}列):')
for r_idx, row in enumerate(t5.rows):
    cells = [cell.text.strip() for cell in row.cells]
    print(f'  行[{r_idx}]: {cells}')

# 表格[6] 行[3]
t6 = doc.tables[6]
print(f'\n表格[6] ({len(t6.rows)}行×{len(t6.columns)}列):')
for r_idx, row in enumerate(t6.rows):
    cells = [cell.text.strip() for cell in row.cells]
    print(f'  行[{r_idx}]: {cells}')

# 表格[7] 行[4]
t7 = doc.tables[7]
print(f'\n表格[7] ({len(t7.rows)}行×{len(t7.columns)}列):')
for r_idx, row in enumerate(t7.rows):
    cells = [cell.text.strip() for cell in row.cells]
    print(f'  行[{r_idx}]: {cells}')
