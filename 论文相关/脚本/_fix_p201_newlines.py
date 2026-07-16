"""修复 [201] 段落换行问题：拆成 1 引导句 + 4 要点段落。

原 [201] 所有内容塞在一个段落里用 ' - ' 分隔，docx 中不会换行。
拆成多段落后，在 [202]（6.2 未来工作）之前插入。
"""
import docx
from docx.oxml.ns import qn
from copy import deepcopy

DOC_PATH = r'论文相关\论文与实验报告\论文1_DL-LNN颤振预测主论文.docx'
doc = docx.Document(DOC_PATH)

# 验证 [201] 和 [202]
p201 = doc.paragraphs[201]
p202 = doc.paragraphs[202]
print(f"[201] style={p201.style.name}")
print(f"[202] style={p202.style.name}, text={p202.text[:40]}")
assert '6.2' in p202.text, f"[202] 应为 6.2 标题，实际: {p202.text[:40]}"
assert '当前实验结论' in p201.text, f"[201] 应为结论段，实际: {p201.text[:40]}"

# 拆分内容
intro_text = '当前实验结论（基于已完成的 Synthetic + Industrial + PHM2010 三个数据集 × 9 模型主实验，v0.4）：'

bullets = [
    '精度优势验证成功：DL-LNN 在 Synthetic（MAE=0.3222, R²=0.9968）与 Industrial（MAE=0.9289, R²=0.9680）两数据集 MAE 排名均跃居 1/9，较次优基线 PINN 分别降低 36.5% 与 2.8%，验证了 Target 归一化机制修复与切屑变薄系数修正后精度优势声明得以恢复。',
    '物理一致性验证成功：DL-LNN 阶段二训练 PCC 达到 0.9953（Synthetic）/ 0.997（Industrial），证明 PCC Loss 确实将预测约束到物理可行域，验证了"结构同构 + 梯度层约束"在塑造物理一致预测函数上的有效性。',
    '跨工况泛化能力初步验证：LOMO/LOCO 协议下 DL-LNN 相比 PINN 平均 MAE 降低 3.3%-4.2%，PCC 提升 7.1%-10.2%，验证了物理约束对未见工况的泛化价值（待完整实验填充）。',
    '诚实分层评估：在 PHM2010 真实数据集（208 样本，目标 std=0.1246 极窄，输入为信号统计量）上 DL-LNN 排名 9/9（MAE=0.1119，末位），树模型（RF/XGBoost）凭借小样本低方差表格数据优势占据前两位；该结果表明 DL-LNN 的优势区间为物理富集数据集，在信号派生数据集上精度落后但保留物理可解释性。',
]

# Step 1: 修改 [201] 为引导句（保留第一个 run 的格式）
first_run = p201.runs[0]
first_run.text = intro_text
for run in p201.runs[1:]:
    run.text = ''
print(f"[201] 已改为引导句: {p201.text[:60]}")

# Step 2: 在 [202] 之前插入 4 个要点段落
# 使用 insert_paragraph_before 方法
for i, bullet in enumerate(bullets):
    new_p = p202.insert_paragraph_before(bullet, style=p201.style)
    print(f"  插入要点 {i+1}: {bullet[:50]}...")

# 保存
doc.save(DOC_PATH)
print(f"\n修复完成，已保存至: {DOC_PATH}")

# 验证
doc2 = docx.Document(DOC_PATH)
print("\n验证（重新读取）：")
for i in range(199, 208):
    if i < len(doc2.paragraphs):
        p = doc2.paragraphs[i]
        print(f"[{i}] ({p.style.name}) {p.text[:70]}")
