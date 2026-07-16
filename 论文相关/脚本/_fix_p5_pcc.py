"""修复 [5] 段落（摘要/引言）的 v0.3 残留 PCC 值。"""
import docx

DOC_PATH = r'论文相关\论文与实验报告\论文1_DL-LNN颤振预测主论文.docx'
doc = docx.Document(DOC_PATH)

p5 = doc.paragraphs[5]
print(f"[5] 修改前: ...{p5.text[p5.text.find('PCC'):p5.text.find('PCC')+120]}...")

# 替换 0.987/0.997 → 0.9953/0.997
full_text = p5.text
new_text = full_text.replace('0.987/0.997', '0.9953/0.997')

if new_text == full_text:
    print("⚠️ 未找到 '0.987/0.997'，尝试其他格式")
    # 尝试其他可能格式
    if '0.987 / 0.997' in full_text:
        new_text = full_text.replace('0.987 / 0.997', '0.9953 / 0.997')
    elif '0.987' in full_text:
        new_text = full_text.replace('0.987', '0.9953')
    else:
        print("❌ [5] 中未找到任何 0.987 变体")

# 应用替换：保留第一个 run 格式
if p5.runs:
    p5.runs[0].text = new_text
    for run in p5.runs[1:]:
        run.text = ''
else:
    p5.add_run(new_text)

doc.save(DOC_PATH)
print(f"\n[5] 修改后: ...{p5.text[p5.text.find('PCC'):p5.text.find('PCC')+120]}...")
print(f"\n保存完成")

# 最终扫描确认
doc2 = docx.Document(DOC_PATH)
remaining = 0
for i, p in enumerate(doc2.paragraphs):
    if '0.987' in p.text:
        remaining += 1
        print(f"残留 [{i}]: {p.text[:80]}")
for t_idx, t in enumerate(doc2.tables):
    for r_idx, row in enumerate(t.rows):
        for c_idx, cell in enumerate(row.cells):
            if '0.987' in cell.text:
                remaining += 1
                print(f"残留 表格[{t_idx}] 行[{r_idx}] 列[{c_idx}]: {cell.text}")

if remaining == 0:
    print("\n✅ 论文1 中已无 0.987 残留，v0.4 一致性完整")
