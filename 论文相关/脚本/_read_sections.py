"""读取 Section 5.2/5.3/6.1 相关段落的完整内容，便于规划重写。"""
import docx

DOC_PATH = r'论文相关\论文与实验报告\论文1_DL-LNN颤振预测主论文.docx'
doc = docx.Document(DOC_PATH)

# 重点段落
target_indices = [151, 152, 153, 154, 155, 156, 157, 158, 159, 199, 200, 201, 202]

for i in target_indices:
    p = doc.paragraphs[i]
    style = p.style.name if p.style else 'None'
    print(f"\n[{i}] ({style}) runs={len(p.runs)}")
    print(f"  文本: {p.text}")

# 检查 [153] 和 [158] 之后是否有表格
print("\n" + "=" * 80)
print("表格检查（Section 5.2/5.3 附近）")
print("=" * 80)
# doc.tables 是文档级表格，但段落与表格的相对位置需要通过 body 元素判断
body = doc.element.body
para_count = 0
table_count = 0
for child in body.iterchildren():
    tag = child.tag.split('}')[-1]
    if tag == 'p':
        # 段落
        text = ''.join(t.text or '' for t in child.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'))
        if 145 <= para_count <= 165:
            print(f"  段落[{para_count}]: {text[:80]}")
        para_count += 1
    elif tag == 'tbl':
        # 表格
        print(f"  >>> 表格[{table_count}] 出现在段落[{para_count}] 之后")
        table_count += 1
