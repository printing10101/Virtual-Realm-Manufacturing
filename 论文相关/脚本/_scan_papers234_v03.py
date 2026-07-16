"""扫描论文2/3/4 中的 v0.3 残留数值。"""
import docx
import os

PAPERS = [
    r'论文相关\论文与实验报告\论文2_PCC_Loss通用化方法论.docx',
    r'论文相关\论文与实验报告\论文3_双分支门控融合架构.docx',
    r'论文相关\论文与实验报告\论文4_CTNN制造应用综述.docx',
]

stale_patterns = ['0.987', '0.3744', '1.2061', '8.89%', '1.60%', '14.8%', '20.000†', 'MAE 尚非最优', '略逊于']

for paper_path in PAPERS:
    if not os.path.exists(paper_path):
        print(f'\n{"="*70}\n文件不存在: {paper_path}\n{"="*70}')
        continue

    doc = docx.Document(paper_path)
    fname = os.path.basename(paper_path)

    print(f'\n{"="*70}')
    print(f'扫描: {fname}')
    print(f'{"="*70}')

    residuals = 0
    # 段落扫描
    for i, p in enumerate(doc.paragraphs):
        for pattern in stale_patterns:
            if pattern in p.text:
                print(f'  [{i}] 段落 "{pattern}": {p.text[:120]}')
                residuals += 1

    # 表格扫描
    for t_idx, t in enumerate(doc.tables):
        for r_idx, row in enumerate(t.rows):
            for c_idx, cell in enumerate(row.cells):
                for pattern in stale_patterns:
                    if pattern in cell.text:
                        print(f'  表格[{t_idx}] 行[{r_idx}] 列[{c_idx}] "{pattern}": {cell.text[:80]}')
                        residuals += 1

    if residuals == 0:
        print(f'  ✅ 无 v0.3 残留')
    else:
        print(f'  ⚠️  发现 {residuals} 处残留')
