"""验证论文1 v0.4 更新结果。"""
import docx

DOC_PATH = r'论文相关\论文与实验报告\论文1_DL-LNN颤振预测主论文.docx'
doc = docx.Document(DOC_PATH)

print("=" * 80)
print("论文1 v0.4 更新验证")
print("=" * 80)

checks = [
    (146, "0.9953", "Section 5.1 PCC"),
    (149, "rank 1/9", "Section 5.1 精度优势"),
    (149, "0.3222", "Section 5.1 Synthetic MAE"),
    (149, "0.9289", "Section 5.1 Industrial MAE"),
    (149, "36.5%", "Section 5.1 PINN 改善幅度"),
    (150, "SOTA", "Section 5.1 双重优势"),
    (150, "双重优势", "Section 5.1 双重优势关键词"),
    (162, "0.9953", "Section 5.4 PCC"),
    (181, "0.9953", "Section 5.x PCC"),
    (201, "1/9", "Section 6.1 结论 rank"),
    (201, "0.9953", "Section 6.1 结论 PCC"),
    (201, "9/9", "Section 6.1 PHM2010 末位"),
]

all_ok = True
for idx, keyword, desc in checks:
    text = doc.paragraphs[idx].text
    found = keyword in text
    status = "✅" if found else "❌"
    if not found:
        all_ok = False
    print(f"{status} [{idx}] {desc}: '{keyword}'")
    if not found:
        # 打印部分内容以便调试
        print(f"    实际内容: {text[:120]}...")

print("=" * 80)
# 检查是否还残留旧的 v0.3 内容
stale_checks = [
    (146, "0.987（Synthetic）", "旧 PCC 值"),
    (149, "略逊于 PINN", "旧精度声明"),
    (150, "设计目标并非在单工况精度上超越", "旧双重优势声明"),
    (162, "PCC=0.987", "旧 PCC 值"),
    (201, "MAE 尚非最优", "旧结论"),
]

print("\n残留旧内容检查：")
for idx, stale_str, desc in stale_checks:
    text = doc.paragraphs[idx].text
    has_stale = stale_str in text
    status = "❌ 仍残留" if has_stale else "✅ 已清除"
    if has_stale:
        all_ok = False
    print(f"{status} [{idx}] {desc}: '{stale_str}'")

print("\n" + "=" * 80)
print(f"总体验证: {'✅ 全部通过' if all_ok else '❌ 存在问题'}")
print("=" * 80)

# 打印更新后的关键段落完整内容
print("\n[149] 更新后完整内容：")
print(doc.paragraphs[149].text)
print("\n[150] 更新后完整内容：")
print(doc.paragraphs[150].text)
print("\n[201] 更新后完整内容：")
print(doc.paragraphs[201].text)
