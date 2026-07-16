"""定位论文中所有 L_pcc 公式相关段落，用于一致性修复。"""
import docx
from pathlib import Path

DOC_PATH = Path(r"c:\Users\Lenovo\Desktop\灵境制造（上线版）\论文相关\论文与实验报告\论文1_DL-LNN颤振预测主论文.docx")

doc = docx.Document(str(DOC_PATH))

print(f"文档总段落数: {len(doc.paragraphs)}")
print("=" * 70)

# 关键字：覆盖 L_pcc 公式的所有可能写法
keywords = [
    "L_pcc", "L_{pcc}", "∂y_pred", "∂y_Tlusty", "梯度一致性",
    "梯度层损失", "∇_x", "Σᵢ", "|∂", "y_pred/∂x",
]

for idx, p in enumerate(doc.paragraphs):
    text = p.text
    if not text.strip():
        continue
    for kw in keywords:
        if kw in text:
            print(f"[P{idx}] {text}")
            print("-" * 70)
            break

print("\n\n=== 表格中也搜索 ===")
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            text = cell.text
            for kw in keywords:
                if kw in text:
                    print(f"[Table{ti} R{ri} C{ci}] {text[:200]}")
                    break
