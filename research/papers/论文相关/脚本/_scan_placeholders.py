"""扫描论文1 中所有"待实验"标记和占位内容，为后续填充做准备。"""
import docx

DOC_PATH = r'论文相关\论文与实验报告\论文1_DL-LNN颤振预测主论文.docx'
doc = docx.Document(DOC_PATH)

print('=' * 70)
print('论文1 占位内容扫描')
print('=' * 70)

# 1. 扫描"待实验"/"待填充"/"占位"标记
print('\n[1] "待实验/待填充/占位" 标记扫描:')
markers = ['待实验', '待填充', '占位', 'TODO', 'TBD', '【待', '【占位']
for i, p in enumerate(doc.paragraphs):
    for marker in markers:
        if marker in p.text:
            print(f'  [{i}] ({marker}): {p.text[:120]}')

# 2. 扫描所有表格，识别占位表格
print('\n[2] 表格概览:')
for t_idx, t in enumerate(doc.tables):
    n_rows = len(t.rows)
    n_cols = len(t.columns)
    # 获取第一行作为标题
    header = [cell.text.strip() for cell in t.rows[0].cells] if n_rows > 0 else []
    print(f'  表格[{t_idx}]: {n_rows}行 × {n_cols}列, 标题: {header[:5]}')

# 3. 扫描 Section 5.2/5.3 相关段落
print('\n[3] Section 5.2/5.3 相关段落:')
for i in range(145, 175):
    if i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        if p.style.name.startswith('Heading') or 'LOMO' in p.text or 'LOCO' in p.text or '消融' in p.text:
            print(f'  [{i}] ({p.style.name}): {p.text[:100]}')
