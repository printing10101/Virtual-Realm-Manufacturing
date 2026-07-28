"""修复论文 P29/P71/P83 的 L_pcc 公式：L1 范数 → L2 范数平方，与代码 losses.py 一致。

代码实现（losses.py 行 155-158）：
    grad_diff = grad_pred - grad_physics
    loss_pcc = torch.mean(grad_diff ** 2)
即 L_pcc = (1/N) · ‖∇_x y_pred − ∇_x y_Tlusty‖² （L2 范数平方，对批次取均值）

原论文公式（L1 范数）：
    P29:  L_pcc = Σᵢ |∂y_pred/∂xᵢ - ∂y_Tlusty/∂xᵢ|
    P71:  L_pcc = (1/N) · Σᵢ ( |∂y_pred/∂xᵢ - ∂y_Tlusty/∂xᵢ| )
    P83:  L_pcc ← λ₃ · Σ_i |g_pred_i - g_ana_i|

修改后（L2 范数平方，与代码一致）：
    P29:  L_pcc = ‖∇_x y_pred − ∇_x y_Tlusty‖²
    P71:  L_pcc = (1/N) · ‖∇_x y_pred − ∇_x y_Tlusty‖²
    P83:  L_pcc ← λ₃ · Σ_i (g_pred_i − g_ana_i)²
"""
import shutil
from pathlib import Path
import docx

DOC_PATH = Path(r"c:\Users\Lenovo\Desktop\灵境制造（上线版）\论文相关\论文与实验报告\论文1_DL-LNN颤振预测主论文.docx")
BACKUP_PATH = DOC_PATH.with_suffix(f".backup_l_pcc_l1_to_l2{DOC_PATH.suffix}")

# 修改前先备份
if not BACKUP_PATH.exists():
    shutil.copy2(str(DOC_PATH), str(BACKUP_PATH))
    print(f"[备份] {BACKUP_PATH.name}")
else:
    print(f"[已存在备份] {BACKUP_PATH.name}")

doc = docx.Document(str(DOC_PATH))

# ============================================================
# P29 Run[12]: 贡献4 中的内联公式
# ============================================================
p29 = doc.paragraphs[29]
old_p29 = p29.runs[12].text
new_p29 = "L_pcc = ‖∇_x y_pred − ∇_x y_Tlusty‖²"
assert old_p29 == "L_pcc = Σᵢ |∂y_pred/∂xᵢ - ∂y_Tlusty/∂xᵢ|", f"P29 Run[12] 文本不匹配: {old_p29!r}"
p29.runs[12].text = new_p29
print(f"[P29 Run[12]] 修改:")
print(f"  旧: {old_p29}")
print(f"  新: {new_p29}")

# ============================================================
# P71 Run[0]: 核心公式定义
# ============================================================
p71 = doc.paragraphs[71]
old_p71 = p71.runs[0].text
new_p71 = "L_pcc = (1/N) · ‖∇_x y_pred − ∇_x y_Tlusty‖²"
assert old_p71 == "L_pcc = (1/N) · Σᵢ ( |∂y_pred/∂xᵢ - ∂y_Tlusty/∂xᵢ| )", f"P71 Run[0] 文本不匹配: {old_p71!r}"
p71.runs[0].text = new_p71
print(f"\n[P71 Run[0]] 修改:")
print(f"  旧: {old_p71}")
print(f"  新: {new_p71}")

# ============================================================
# P83 Run[73]: 算法伪代码第 20 行
# ============================================================
p83 = doc.paragraphs[83]
old_p83 = p83.runs[73].text
new_p83 = "20:         L_pcc ← λ₃ · Σ_i (g_pred_i − g_ana_i)²"
assert old_p83 == "20:         L_pcc ← λ₃ · Σ_i |g_pred_i - g_ana_i|", f"P83 Run[73] 文本不匹配: {old_p83!r}"
p83.runs[73].text = new_p83
print(f"\n[P83 Run[73]] 修改:")
print(f"  旧: {old_p83}")
print(f"  新: {new_p83}")

# ============================================================
# 保存
# ============================================================
doc.save(str(DOC_PATH))
print(f"\n[保存] {DOC_PATH.name}")

# ============================================================
# 验证：重新打开文档检查修改是否生效
# ============================================================
print("\n=== 验证修改 ===")
doc2 = docx.Document(str(DOC_PATH))
for idx, expected in [(29, new_p29), (71, new_p71), (83, new_p83)]:
    actual = doc2.paragraphs[idx].text
    if idx == 29:
        # P29 是整段文本，只检查包含新公式
        ok = new_p29 in actual
        print(f"P{idx}: {'✓' if ok else '✗'} 包含新公式")
        if not ok:
            print(f"  实际: {actual}")
    elif idx == 71:
        ok = actual == new_p71
        print(f"P{idx}: {'✓' if ok else '✗'}")
        if not ok:
            print(f"  期望: {new_p71}")
            print(f"  实际: {actual}")
    elif idx == 83:
        # P83 是多行伪代码，检查第 20 行
        ok = new_p83 in actual
        print(f"P{idx}: {'✓' if ok else '✗'} 包含新伪代码")
        if not ok:
            # 找到第 20 行
            for line in actual.split("\n"):
                if line.strip().startswith("20:"):
                    print(f"  实际第20行: {line}")
                    break
