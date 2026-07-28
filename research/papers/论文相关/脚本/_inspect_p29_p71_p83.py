"""检查 P29, P71, P83 的 run 结构，准备修改 L_pcc 公式。"""
import docx
from pathlib import Path

DOC_PATH = Path(r"c:\Users\Lenovo\Desktop\灵境制造（上线版）\论文相关\论文与实验报告\论文1_DL-LNN颤振预测主论文.docx")

doc = docx.Document(str(DOC_PATH))

TARGET_INDICES = [29, 71, 83]

for idx in TARGET_INDICES:
    if idx >= len(doc.paragraphs):
        print(f"[P{idx}] 越界")
        continue
    p = doc.paragraphs[idx]
    print(f"\n=== P{idx} ===")
    print(f"文本: {p.text}")
    print(f"Run 数量: {len(p.runs)}")
    for ri, run in enumerate(p.runs):
        font = run.font
        italic = font.italic
        bold = font.bold
        size = font.size
        name = font.name
        print(f"  Run[{ri}]: italic={italic}, bold={bold}, size={size}, name={name}")
        print(f"    text={repr(run.text)}")
