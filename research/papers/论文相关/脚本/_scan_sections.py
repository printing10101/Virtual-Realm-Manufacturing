"""扫描论文1 章节结构，定位 Section 5.2/5.3 占位内容，并修复 [201] 换行符。"""
import docx
from docx.oxml.ns import qn

DOC_PATH = r'论文相关\论文与实验报告\论文1_DL-LNN颤振预测主论文.docx'
doc = docx.Document(DOC_PATH)

print("=" * 80)
print("论文1 章节结构扫描")
print("=" * 80)

# 扫描所有标题段落（包含"5."或"第 5"或"LOMO"或"消融"关键词）
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    # 识别标题
    if (text.startswith('5.') or text.startswith('6.') or
        text.startswith('第 5') or text.startswith('第 6') or
        'LOMO' in text or '消融' in text or
        text.startswith('4.8') or text.startswith('4.6') or text.startswith('4.7')):
        style = p.style.name if p.style else 'None'
        print(f"[{i}] ({style}) {text[:100]}")

print("\n" + "=" * 80)
print("[201] 段落 XML 检查（检查换行符渲染）")
print("=" * 80)
p201 = doc.paragraphs[201]
print(f"段落数据: {repr(p201.text[:200])}")
print(f"Run 数量: {len(p201.runs)}")
for i, run in enumerate(p201.runs[:5]):
    has_break = run._element.findall(qn('w:br'))
    print(f"  Run[{i}]: break={len(has_break)}, text={repr(run.text[:80])}")
