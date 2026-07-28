"""扫描论文1 中所有 v0.3 残留数值，确保 v0.4 一致性。

重点检查：
- 0.987（旧 PCC 值）
- "略逊于"/"非最优"/"落后"（旧精度声明）
- 14.8%/4.6%/8.89%/1.60%（旧差距百分比）
- 0.3744/1.2061（旧 MAE 值）
"""
import docx

DOC_PATH = r'论文相关\论文与实验报告\论文1_DL-LNN颤振预测主论文.docx'
doc = docx.Document(DOC_PATH)

# v0.3 残留关键词
stale_patterns = [
    '0.987',
    '略逊于',
    '非最优',
    '14.8%',
    '4.6%',
    '8.89%',
    '1.60%',
    '0.3744',
    '1.2061',
    '设计目标并非',
    'MAE 尚非最优',
]

print("=" * 80)
print("论文1 v0.3 残留数值扫描")
print("=" * 80)

found_any = False
for i, p in enumerate(doc.paragraphs):
    text = p.text
    if not text:
        continue
    for pattern in stale_patterns:
        if pattern in text:
            found_any = True
            # 上下文
            idx = text.find(pattern)
            start = max(0, idx - 30)
            end = min(len(text), idx + len(pattern) + 30)
            context = text[start:end]
            print(f"[{i}] '{pattern}': ...{context}...")

if not found_any:
    print("✅ 未发现 v0.3 残留数值，论文1 v0.4 一致性良好")

# 额外检查：表格中的数值
print("\n" + "=" * 80)
print("表格内容检查（查找 0.987 等旧值）")
print("=" * 80)
for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            text = cell.text
            for pattern in ['0.987', '0.3744', '1.2061', '14.8%', '8.89%']:
                if pattern in text:
                    print(f"表格[{t_idx}] 行[{r_idx}] 列[{c_idx}]: '{text}' 含 '{pattern}'")

print("\n扫描完成")
