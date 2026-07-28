"""更新表格[5]（表 2 主实验表）为 v0.4 数据。

v0.4 数据来源：python/experiments/results/all_experiments_results.json
- 移除 GP 的 † 发散标记（v0.4 已通过 Optuna 修复）
- 更新所有模型的 Synthetic/Industrial MAE
- 重新计算"已完成列平均 MAE"= (Synthetic + Industrial) / 2
- 更新 DL-LNN 的 PCC: 0.987/0.997 → 0.9953/0.997
"""
import docx

DOC_PATH = r'论文相关\论文与实验报告\论文1_DL-LNN颤振预测主论文.docx'
doc = docx.Document(DOC_PATH)

# v0.4 数据（来自 all_experiments_results.json）
v04_data = {
    1: {"name": "SVR",            "syn": 2.1332, "ind": 1.3029},
    2: {"name": "Random Forest",  "syn": 1.3528, "ind": 1.0576},
    3: {"name": "XGBoost",        "syn": 1.2144, "ind": 1.1051},
    4: {"name": "BPNN",           "syn": 0.5231, "ind": 1.2225},
    5: {"name": "LSTM",           "syn": 0.5663, "ind": 0.9496},
    6: {"name": "Transformer",    "syn": 1.1246, "ind": 6.3370},
    7: {"name": "PINN",           "syn": 0.5076, "ind": 0.9560},
    8: {"name": "Gaussian Process", "syn": 2.6367, "ind": 2.4488},
    9: {"name": "DL-LNN（本文）",  "syn": 0.3222, "ind": 0.9289, "pcc": "0.9953 / 0.997"},
}

table = doc.tables[5]

def set_cell_text(cell, new_text):
    """设置单元格文本，保留第一个 run 的格式。"""
    # 清除所有现有段落（保留第一个）
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = new_text
        for run in p.runs[1:]:
            run.text = ''
    else:
        p.add_run(new_text)

# 更新每一行
for row_idx, data in v04_data.items():
    row = table.rows[row_idx]
    avg = (data["syn"] + data["ind"]) / 2

    # 列[5] Synthetic
    set_cell_text(row.cells[5], f'{data["syn"]:.4f}')
    # 列[6] Industrial
    set_cell_text(row.cells[6], f'{data["ind"]:.4f}')
    # 列[7] 已完成列平均 MAE
    set_cell_text(row.cells[7], f'{avg:.4f}')
    # 列[8] PCC（仅 DL-LNN）
    if "pcc" in data:
        set_cell_text(row.cells[8], data["pcc"])

    print(f"行[{row_idx}] {data['name']}: Synthetic={data['syn']:.4f}, Industrial={data['ind']:.4f}, avg={avg:.4f}")

doc.save(DOC_PATH)
print(f"\n表格[5] v0.4 更新完成，已保存至: {DOC_PATH}")

# 验证
doc2 = docx.Document(DOC_PATH)
t = doc2.tables[5]
print("\n验证（DL-LNN 行）：")
for c_idx, cell in enumerate(t.rows[9].cells):
    print(f"  列[{c_idx}]: {cell.text}")
print("\n验证（GP 行）：")
for c_idx, cell in enumerate(t.rows[8].cells):
    print(f"  列[{c_idx}]: {cell.text}")
