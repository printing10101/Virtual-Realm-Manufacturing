"""修复 [155] 和 [159] 段落换行问题：拆成 1 引导句 + 多个要点段落。

与 [201] 修复方式相同：使用 insert_paragraph_before() 在下一段前插入要点段落。
"""
import docx

DOC_PATH = r'论文相关\论文与实验报告\论文1_DL-LNN颤振预测主论文.docx'
doc = docx.Document(DOC_PATH)


def split_paragraph(doc, idx, intro_text, bullets):
    """将 [idx] 段落拆成引导句 + 多个要点段落。

    要点段落插入到 [idx+1] 之前（即原 [idx] 之后）。
    """
    p = doc.paragraphs[idx]
    p_next = doc.paragraphs[idx + 1]

    # Step 1: 修改 [idx] 为引导句
    first_run = p.runs[0]
    first_run.text = intro_text
    for run in p.runs[1:]:
        run.text = ''

    # Step 2: 在 [idx+1] 之前插入要点段落
    for bullet in bullets:
        p_next.insert_paragraph_before(bullet, style=p.style)

    print(f'[{idx}] 拆分完成：1 引导句 + {len(bullets)} 要点段落')


# ============================================================
# 修复 [155]: LOMO/LOCO 结果分析
# ============================================================
p155_intro = '结果分析：'
p155_bullets = [
    'DL-LNN 在跨材料场景下平均 MAE 比 PINN 降低 3.3%，PCC 提升 10.2%（0.943 vs 0.856）。',
    'DL-LNN 在跨工况场景下平均 MAE 比 PINN 降低 4.2%，PCC 提升 7.1%（0.918 vs 0.857）。',
    '外推到第 5 种材料时，纯数据驱动方法（Transformer、BPNN）MAE 显著上升，而 DL-LNN 由于 PCC Loss 约束预测维持在物理可行域，泛化能力更稳健。',
    '验证了 PCC Loss 在跨工况泛化上的核心价值：物理约束使模型在未见工况上仍能保持物理合理性。',
]
split_paragraph(doc, 155, p155_intro, p155_bullets)

# 修复 [155] 后，原 [159] 现在变为 [159 + 4 = 163]
# 需要重新读取文档以获取新索引
doc.save(DOC_PATH)
print('\n[155] 修复已保存，重新加载文档以修复 [159]...')

doc = docx.Document(DOC_PATH)

# ============================================================
# 修复原 [159]（现在索引变为 163）: 消融实验关键发现
# ============================================================
# 先验证 [163] 是否为消融分析段落
p163 = doc.paragraphs[163]
print(f'\n[163] 验证: {p163.text[:80]}')
assert '关键发现' in p163.text, f'[163] 不是消融分析段落: {p163.text[:80]}'

p163_intro = '关键发现：'
p163_bullets = [
    'A1 vs Full：去除 L_phys 后 MAE 上升 11.7%，PCC 下降 4.9%，证明数值层物理约束提供物理硬边界。',
    'A2 vs Full：去除 L_pcc 后 MAE 反而下降 4.4%，但 PCC 显著下降至 0.892（-10.5%），证明梯度层一致性损失在精度上有少量代价但显著提升物理合理性——这是 DL-LNN 与 PINN 的核心差异化竞争力。',
    'A3 vs Full：去除两阶段训练后 MAE 上升 23.5%，证明解析预训练对小样本冷启动至关重要。',
    'A4 敏感性：λ₃=0.10 最优；过小（0.01）物理约束不足，过大（1.0）过度偏向物理损失而牺牲数据拟合。',
    'A6 门控策略：固定 α=0（纯解析）精度最差，固定 α=1（纯数据）PCC 最低，自适应 α(x) 在精度与物理一致性间取得最佳平衡。',
    'A7 主干对比：LTC 优于 CNN（5.6%）优于 MLP（10.9%），证明连续时间 ODE 结构相比离散时间网络的结构性优势。',
]
split_paragraph(doc, 163, p163_intro, p163_bullets)

doc.save(DOC_PATH)
print(f'\n论文1 [155]/[159] 换行修复完成，已保存至: {DOC_PATH}')

# 验证
doc2 = docx.Document(DOC_PATH)
print('\n=== 验证 ===')
for i in [155, 156, 157, 158, 159, 160, 161, 162, 163, 164, 165, 166, 167, 168, 169, 170]:
    if i < len(doc2.paragraphs):
        p = doc2.paragraphs[i]
        text = p.text
        print(f'[{i}] ({p.style.name}): {text[:100]}')
