"""修复论文1 残留的 v0.3 数值：[5] 段落 + 表格[5] 行[9]。"""
import docx

DOC_PATH = r'论文相关\论文与实验报告\论文1_DL-LNN颤振预测主论文.docx'
doc = docx.Document(DOC_PATH)

# 检查 [5] 段落
p5 = doc.paragraphs[5]
print(f"[5] 完整内容：")
print(p5.text)
print(f"\n[5] runs 数量: {len(p5.runs)}")

# 检查表格[5]
table5 = doc.tables[5]
print(f"\n表格[5] 行数: {len(table5.rows)}, 列数: {len(table5.columns)}")
print("表格[5] 第 9 行内容：")
for c_idx, cell in enumerate(table5.rows[9].cells):
    print(f"  列[{c_idx}]: {cell.text}")

# 打印表格[5] 的标题行（第 0 行）以理解结构
print("\n表格[5] 标题行（第 0 行）：")
for c_idx, cell in enumerate(table5.rows[0].cells):
    print(f"  列[{c_idx}]: {cell.text}")

# 打印前几行
print("\n表格[5] 第 1-3 行：")
for r_idx in range(1, min(4, len(table5.rows))):
    print(f"  行[{r_idx}]: ", end='')
    for c_idx, cell in enumerate(table5.rows[r_idx].cells):
        print(f"[{c_idx}]={cell.text[:20]}", end=' ')
    print()
