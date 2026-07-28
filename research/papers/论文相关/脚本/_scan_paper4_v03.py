"""扫描论文4 的 v0.3 残留。"""
import docx
import os

DOC_PATH = r'论文相关\论文与实验报告\论文4_连续时间神经网络制造应用综述.docx'

if not os.path.exists(DOC_PATH):
    print(f'文件不存在: {DOC_PATH}')
else:
    doc = docx.Document(DOC_PATH)
    stale_patterns = ['0.987', '0.3744', '1.2061', '8.89%', '1.60%', '14.8%', '20.000†', 'MAE 尚非最优', '略逊于']

    print('=' * 70)
    print('论文4 v0.3 残留扫描')
    print('=' * 70)

    residuals = 0
    for i, p in enumerate(doc.paragraphs):
        for pattern in stale_patterns:
            if pattern in p.text:
                print(f'  [{i}] 段落 "{pattern}": {p.text[:120]}')
                residuals += 1

    for t_idx, t in enumerate(doc.tables):
        for r_idx, row in enumerate(t.rows):
            for c_idx, cell in enumerate(row.cells):
                for pattern in stale_patterns:
                    if pattern in cell.text:
                        print(f'  表格[{t_idx}] 行[{r_idx}] 列[{c_idx}] "{pattern}": {cell.text[:80]}')
                        residuals += 1

    if residuals == 0:
        print('  ✅ 无 v0.3 残留')
    else:
        print(f'  ⚠️  发现 {residuals} 处残留')
