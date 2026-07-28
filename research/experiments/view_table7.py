"""
查看表7的详细内容
"""
from docx import Document
from pathlib import Path

paper_path = Path("../../docs/DL-LNN-论文-更新版.docx")
doc = Document(str(paper_path))

print("=" * 80)
print("表7详细内容")
print("=" * 80)

# 表7是第7个表格（索引6）
if len(doc.tables) >= 7:
    table7 = doc.tables[6]
    
    print(f"\n表格行数: {len(table7.rows)}")
    print(f"表格列数: {len(table7.columns)}")
    
    print("\n表格内容:")
    for i, row in enumerate(table7.rows):
        print(f"\n行 {i}:")
        for j, cell in enumerate(row.cells):
            text = cell.text.strip()
            if text:
                print(f"  列{j}: {text[:100]}")
else:
    print("表格数量不足7个")
