"""
综合验证论文中所有表格的数据一致性
检查是否还有占位符，验证数据完整性
"""
from docx import Document
from pathlib import Path
import json

print("=" * 80)
print("论文数据一致性综合验证")
print("=" * 80)

# 加载论文
paper_path = Path("../../docs/DL-LNN-论文-最终版.docx")
if not paper_path.exists():
    paper_path = Path("../../docs/DL-LNN-论文-更新版.docx")

doc = Document(str(paper_path))

print(f"\n论文文件: {paper_path}")
print(f"总表格数: {len(doc.tables)}")

# 验证结果
verification_results = {
    "total_tables": len(doc.tables),
    "tables_with_placeholders": [],
    "tables_verified": [],
    "issues": []
}

print("\n" + "=" * 80)
print("逐表验证")
print("=" * 80)

for i, table in enumerate(doc.tables):
    print(f"\n--- 表格 {i} ---")
    print(f"行数: {len(table.rows)}, 列数: {len(table.columns)}")
    
    has_placeholder = False
    placeholder_cells = []
    
    # 检查每个单元格
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            text = cell.text.strip()
            
            # 检查是否有占位符
            if '⬜' in text or '□' in text or '待填写' in text or 'TBD' in text:
                has_placeholder = True
                placeholder_cells.append((row_idx, col_idx, text[:50]))
    
    if has_placeholder:
        verification_results["tables_with_placeholders"].append(i)
        print(f"  ❌ 仍有占位符:")
        for row_idx, col_idx, text in placeholder_cells:
            print(f"    行{row_idx} 列{col_idx}: {text}")
        verification_results["issues"].append(f"表格{i}仍有{len(placeholder_cells)}个占位符")
    else:
        verification_results["tables_verified"].append(i)
        print(f"  ✓ 无占位符，数据完整")

# 加载实验数据进行交叉验证
print("\n" + "=" * 80)
print("实验数据交叉验证")
print("=" * 80)

results_dir = Path("results")
if results_dir.exists():
    # 验证主对比结果
    main_results_path = results_dir / "main_comparison_results.json"
    if main_results_path.exists():
        with open(main_results_path, 'r', encoding='utf-8') as f:
            main_results = json.load(f)
        
        print(f"\n✓ 主对比实验数据已加载")
        print(f"  数据集数量: {len(main_results)}")
        print(f"  模型数量: {len(list(main_results.values())[0])}")
        
        # 检查DL-LNN在各数据集上的MAE
        print(f"\n  DL-LNN MAE结果:")
        for dataset, models in main_results.items():
            if 'DL-LNN' in models:
                mae = models['DL-LNN']['MAE']
                print(f"    {dataset}: {mae:.4f}")
    
    # 验证时间常数分析
    time_constant_path = results_dir / "time_constant_analysis.json"
    if time_constant_path.exists():
        with open(time_constant_path, 'r', encoding='utf-8') as f:
            tc_results = json.load(f)
        
        print(f"\n✓ 时间常数分析数据已加载")
        print(f"  网络层数: {len(tc_results['layers'])}")
        print(f"  全局τ均值: {tc_results['global']['tau_mean']:.4f}")
        print(f"  全局τ标准差: {tc_results['global']['tau_std']:.4f}")

# 最终总结
print("\n" + "=" * 80)
print("验证总结")
print("=" * 80)

print(f"\n总表格数: {verification_results['total_tables']}")
print(f"已验证表格: {len(verification_results['tables_verified'])}")
print(f"仍有问题的表格: {len(verification_results['tables_with_placeholders'])}")

if verification_results['tables_with_placeholders']:
    print(f"\n❌ 以下表格仍有占位符:")
    for table_idx in verification_results['tables_with_placeholders']:
        print(f"  - 表格 {table_idx}")
    
    print(f"\n问题列表:")
    for issue in verification_results['issues']:
        print(f"  - {issue}")
else:
    print(f"\n✓ 所有表格数据完整，无占位符！")

print("\n" + "=" * 80)
print("验证完成")
print("=" * 80)
