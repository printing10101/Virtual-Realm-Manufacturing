"""
验证论文最终版数据完整性
"""
from pathlib import Path
from docx import Document


def verify_paper():
    """验证论文中的图片和数据"""
    
    paper_path = Path("../../docs/DL-LNN-论文-最终版-完整.docx")
    if not paper_path.exists():
        print(f"错误: 论文文件不存在: {paper_path}")
        return
    
    print("=" * 80)
    print("验证论文最终版数据完整性")
    print("=" * 80)
    
    doc = Document(str(paper_path))
    
    # 统计图片
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_count += 1
    
    print(f"\n✓ 文档中的图片数量: {image_count}")
    
    # 统计表格
    table_count = len(doc.tables)
    print(f"✓ 文档中的表格数量: {table_count}")
    
    # 检查关键占位符
    placeholder_count = 0
    for para in doc.paragraphs:
        text = para.text
        if '⬜' in text or '待补充' in text or '待更新' in text:
            placeholder_count += 1
            print(f"  ⚠ 发现占位符: {text[:50]}...")
    
    print(f"\n✓ 关键占位符数量: {placeholder_count}")
    
    # 总结
    print("\n" + "=" * 80)
    print("验证完成")
    print("=" * 80)
    
    if image_count >= 2 and placeholder_count == 0:
        print("✓ 论文数据完整性验证通过")
        print(f"  - 图片: {image_count} 张")
        print(f"  - 表格: {table_count} 个")
        print(f"  - 占位符: {placeholder_count} 个")
    else:
        print("⚠ 论文仍存在问题")
        if image_count < 2:
            print(f"  - 图片数量不足: {image_count}/2")
        if placeholder_count > 0:
            print(f"  - 仍有占位符: {placeholder_count} 个")


if __name__ == "__main__":
    verify_paper()
