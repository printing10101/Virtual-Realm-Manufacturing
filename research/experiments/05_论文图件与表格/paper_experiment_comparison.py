#!/usr/bin/env python3
"""
论文与实验结果对比检查脚本
"""

import json
import os
from pathlib import Path
from docx import Document


def extract_paper_content(docx_path):
    """提取论文内容"""
    try:
        doc = Document(docx_path)
        content = {"paragraphs": [], "tables": [], "full_text": ""}

        # 提取段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                content["paragraphs"].append(text)
                content["full_text"] += text + "\n"

        # 提取表格
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            content["tables"].append({"index": i, "data": table_data})

        return content
    except Exception as e:
        print(f"提取论文内容时出错: {e}")
        return None


def check_json_results(results_dir):
    """检查JSON结果文件"""
    json_files = {
        "main_comparison": "main_comparison_results.json",
        "cross_condition": "cross_condition_results.json",
        "ablation": "ablation_results.json",
        "active_learning": "active_learning_results.json",
        "time_constant": "time_constant_analysis.json",
    }

    results = {}
    for name, filename in json_files.items():
        filepath = os.path.join(results_dir, filename)
        if os.path.exists(filepath):
            with open(filepath, "r", encoding="utf-8") as f:
                results[name] = json.load(f)
        else:
            results[name] = None

    return results


def check_figures(figures_dir):
    """检查图表文件"""
    expected_figures = [
        "main_results_synthetic.png",
        "main_results_industrial.png",
        "lomo_heatmap.png",
        "loco_heatmap.png",
        "ablation_study.png",
    ]

    figures_status = {}
    for fig in expected_figures:
        filepath = os.path.join(figures_dir, fig)
        figures_status[fig] = os.path.exists(filepath)

    return figures_status


def check_table_files(results_dir):
    """检查表格数据文件"""
    table_files = {
        "table1": ["table1_synthetic_results.csv", "table1_synthetic_results.tex"],
        "table2": ["table2_industrial_results.csv", "table2_industrial_results.tex"],
        "table3": ["table3_lomo_average.csv"],
        "table4": ["table4_loco_average.csv"],
        "table5": ["table5_ablation_results.csv"],
    }

    table_status = {}
    for table_name, files in table_files.items():
        table_status[table_name] = {}
        for file in files:
            filepath = os.path.join(results_dir, file)
            table_status[table_name][file] = os.path.exists(filepath)

    return table_status


def analyze_main_comparison(results):
    """分析主对比实验结果"""
    if not results.get("main_comparison"):
        return None

    data = results["main_comparison"]
    analysis = {"datasets": list(data.keys()), "models": [], "metrics": [], "completeness": {}}

    # 检查每个数据集的模型和指标
    for dataset_name, dataset_data in data.items():
        models = list(dataset_data.keys())
        if models:
            analysis["models"] = models
            metrics = list(dataset_data[models[0]].keys())
            analysis["metrics"] = metrics
            analysis["completeness"][dataset_name] = {"num_models": len(models), "num_metrics": len(metrics)}

    return analysis


def analyze_cross_condition(results):
    """分析跨条件泛化实验结果"""
    if not results.get("cross_condition"):
        return None

    data = results["cross_condition"]
    analysis = {
        "LOMO": {"materials": list(data.get("LOMO", {}).keys()), "models": [], "complete": False},
        "LOCO": {"conditions": list(data.get("LOCO", {}).keys()), "models": [], "complete": False},
    }

    # 检查LOMO
    if "LOMO" in data:
        lomo_data = data["LOMO"]
        if "Average" in lomo_data:
            analysis["LOMO"]["models"] = list(lomo_data["Average"].keys())
            analysis["LOMO"]["complete"] = len(analysis["LOMO"]["models"]) > 0

    # 检查LOCO
    if "LOCO" in data:
        loco_data = data["LOCO"]
        if "Average" in loco_data:
            analysis["LOCO"]["models"] = list(loco_data["Average"].keys())
            analysis["LOCO"]["complete"] = len(analysis["LOCO"]["models"]) > 0

    return analysis


def analyze_ablation(results):
    """分析消融实验结果"""
    if not results.get("ablation"):
        return None

    data = results["ablation"]
    if "ablation" in data:
        ablation_data = data["ablation"]
        return {
            "configurations": list(ablation_data.keys()),
            "num_configurations": len(ablation_data),
            "complete": len(ablation_data) >= 5,  # Full Model + 4 variants
        }

    return None


def analyze_time_constant(results):
    """分析时间常数分析结果"""
    if not results.get("time_constant"):
        return None

    data = results["time_constant"]
    analysis = {
        "num_layers": len(data.get("layers", [])),
        "has_global": "global" in data,
        "global_stats": {},
        "complete": False,
    }

    if "global" in data:
        global_stats = data["global"]
        analysis["global_stats"] = {
            "tau_mean": global_stats.get("tau_mean"),
            "tau_std": global_stats.get("tau_std"),
            "tau_min": global_stats.get("tau_min"),
            "tau_max": global_stats.get("tau_max"),
            "tau_median": global_stats.get("tau_median"),
        }
        analysis["complete"] = analysis["num_layers"] >= 3 and analysis["has_global"]

    return analysis


def generate_report(paper_content, json_results, figures_status, table_status):
    """生成对比报告"""
    report = []
    report.append("=" * 80)
    report.append("论文与实验结果对比检查报告")
    report.append("=" * 80)
    report.append("")

    # 1. 实验完成情况
    report.append("一、实验完成情况")
    report.append("-" * 80)

    # 主对比实验
    main_analysis = analyze_main_comparison(json_results)
    if main_analysis:
        report.append(f"✓ 主对比实验: 已完成")
        report.append(f"  - 数据集数量: {len(main_analysis['datasets'])}")
        report.append(f"  - 数据集: {', '.join(main_analysis['datasets'])}")
        report.append(f"  - 模型数量: {len(main_analysis['models'])}")
        report.append(f"  - 模型: {', '.join(main_analysis['models'])}")
        report.append(f"  - 评估指标: {', '.join(main_analysis['metrics'])}")
    else:
        report.append(f"✗ 主对比实验: 未完成或数据缺失")
    report.append("")

    # 跨条件泛化实验
    cross_analysis = analyze_cross_condition(json_results)
    if cross_analysis:
        report.append(f"✓ 跨条件泛化实验: 已完成")
        report.append(f"  - LOMO (Leave-One-Material-Out):")
        report.append(f"    · 材料数量: {len(cross_analysis['LOMO']['materials'])}")
        report.append(f"    · 材料: {', '.join(cross_analysis['LOMO']['materials'][:5])}...")
        report.append(f"    · 模型: {', '.join(cross_analysis['LOMO']['models'])}")
        report.append(f"  - LOCO (Leave-One-Condition-Out):")
        report.append(f"    · 条件数量: {len(cross_analysis['LOCO']['conditions'])}")
        report.append(f"    · 模型: {', '.join(cross_analysis['LOCO']['models'])}")
    else:
        report.append(f"✗ 跨条件泛化实验: 未完成或数据缺失")
    report.append("")

    # 消融实验
    ablation_analysis = analyze_ablation(json_results)
    if ablation_analysis:
        report.append(f"✓ 消融实验: 已完成")
        report.append(f"  - 配置数量: {ablation_analysis['num_configurations']}")
        report.append(f"  - 配置: {', '.join(ablation_analysis['configurations'])}")
    else:
        report.append(f"✗ 消融实验: 未完成或数据缺失")
    report.append("")

    # 时间常数分析
    time_analysis = analyze_time_constant(json_results)
    if time_analysis:
        report.append(f"✓ 时间常数分析: 已完成")
        report.append(f"  - 网络层数: {time_analysis['num_layers']}")
        report.append(
            f"  - 全局统计: τ_mean={time_analysis['global_stats'].get('tau_mean', 'N/A'):.4f}, "
            f"τ_std={time_analysis['global_stats'].get('tau_std', 'N/A'):.4f}"
        )
    else:
        report.append(f"✗ 时间常数分析: 未完成或数据缺失")
    report.append("")

    # 主动学习实验
    if json_results.get("active_learning"):
        al_data = json_results["active_learning"]
        report.append(f"✓ 主动学习实验: 已完成")
        report.append(f"  - 数据比例点: {len(al_data.get('active_learning', []))}")
        report.append(f"  - 包含随机基线对比: {'是' if 'random_baseline' in al_data else '否'}")
    else:
        report.append(f"✗ 主动学习实验: 未完成或数据缺失")
    report.append("")

    # 2. 图表生成情况
    report.append("二、图表生成情况")
    report.append("-" * 80)

    report.append("可视化图表:")
    for fig_name, exists in figures_status.items():
        status = "✓" if exists else "✗"
        report.append(f"  {status} {fig_name}")
    report.append("")

    report.append("表格数据文件:")
    for table_name, files in table_status.items():
        report.append(f"  {table_name}:")
        for file_name, exists in files.items():
            status = "✓" if exists else "✗"
            report.append(f"    {status} {file_name}")
    report.append("")

    # 3. 数据一致性检查
    report.append("三、数据一致性检查")
    report.append("-" * 80)

    # 检查主对比实验数据
    if main_analysis:
        report.append("主对比实验数据检查:")
        for dataset_name in main_analysis["datasets"]:
            dataset_data = json_results["main_comparison"][dataset_name]
            num_models = len(dataset_data)
            num_metrics = len(list(dataset_data.values())[0]) if num_models > 0 else 0
            report.append(f"  ✓ {dataset_name}: {num_models}个模型, {num_metrics}个指标")
    report.append("")

    # 检查时间常数数据
    if time_analysis and time_analysis["complete"]:
        report.append("时间常数数据检查:")
        for layer_data in json_results["time_constant"]["layers"]:
            layer_num = layer_data["layer"]
            tau_mean = layer_data["tau_mean"]
            tau_values_count = len(layer_data["tau_values"])
            report.append(f"  ✓ Layer {layer_num}: τ_mean={tau_mean:.4f}, {tau_values_count}个τ值")
        report.append(f"  ✓ Global: τ_mean={time_analysis['global_stats']['tau_mean']:.4f}")
    report.append("")

    # 4. 问题与建议
    report.append("四、问题与建议")
    report.append("-" * 80)

    issues = []

    # 检查缺失的图表
    missing_figures = [fig for fig, exists in figures_status.items() if not exists]
    if missing_figures:
        issues.append(f"缺失图表: {', '.join(missing_figures)}")

    # 检查缺失的表格文件
    missing_tables = []
    for table_name, files in table_status.items():
        for file_name, exists in files.items():
            if not exists:
                missing_tables.append(file_name)
    if missing_tables:
        issues.append(f"缺失表格文件: {', '.join(missing_tables)}")

    # 检查实验完整性
    if not main_analysis:
        issues.append("主对比实验数据缺失")
    if not cross_analysis:
        issues.append("跨条件泛化实验数据缺失")
    if not ablation_analysis:
        issues.append("消融实验数据缺失")
    if not time_analysis:
        issues.append("时间常数分析数据缺失")

    if issues:
        for i, issue in enumerate(issues, 1):
            report.append(f"{i}. {issue}")
    else:
        report.append("✓ 未发现明显问题")

    report.append("")
    report.append("=" * 80)
    report.append("报告生成完成")
    report.append("=" * 80)

    return "\n".join(report)


def main():
    # 设置路径 - 使用绝对路径
    base_dir = Path(r"c:\Users\Lenovo\Desktop\灵境制造（上线版）")
    paper_path = base_dir / "docs" / "DL-LNN-论文-最终版.docx"
    results_dir = base_dir / "python" / "experiments" / "results"
    figures_dir = results_dir / "figures"

    print("开始检查论文与实验结果...")
    print("")

    # 1. 提取论文内容
    print("1. 提取论文内容...")
    paper_content = extract_paper_content(paper_path)
    if paper_content:
        print(f"   ✓ 成功提取 {len(paper_content['paragraphs'])} 个段落")
        print(f"   ✓ 成功提取 {len(paper_content['tables'])} 个表格")
    else:
        print("   ✗ 提取论文内容失败")
    print("")

    # 2. 检查JSON结果
    print("2. 检查实验结果JSON文件...")
    json_results = check_json_results(results_dir)
    for name, data in json_results.items():
        status = "✓" if data else "✗"
        print(f"   {status} {name}")
    print("")

    # 3. 检查图表
    print("3. 检查可视化图表...")
    figures_status = check_figures(figures_dir)
    for fig, exists in figures_status.items():
        status = "✓" if exists else "✗"
        print(f"   {status} {fig}")
    print("")

    # 4. 检查表格文件
    print("4. 检查表格数据文件...")
    table_status = check_table_files(results_dir)
    for table_name, files in table_status.items():
        print(f"   {table_name}:")
        for file_name, exists in files.items():
            status = "✓" if exists else "✗"
            print(f"     {status} {file_name}")
    print("")

    # 5. 生成详细报告
    print("5. 生成详细对比报告...")
    report = generate_report(paper_content, json_results, figures_status, table_status)

    # 保存报告
    report_path = results_dir / "comparison_report.txt"
    with open(report_path, "w", encoding="utf-8") as f:
        f.write(report)

    print(f"   ✓ 报告已保存到: {report_path}")
    print("")
    print(report)


if __name__ == "__main__":
    main()
