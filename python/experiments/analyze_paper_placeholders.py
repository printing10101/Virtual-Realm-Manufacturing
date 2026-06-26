"""
分析论文最终版文档，找出所有占位符
"""
import re
from pathlib import Path
from docx import Document

def analyze_paper():
    """分析论文中的占位符"""
    paper_path = Path("../../docs/DL-LNN-论文-最终版.docx")
    
    if not paper_path.exists():
        print(f"错误: 论文文件不存在: {paper_path}")
        return
    
    print("=" * 80)
    print("分析论文最终版文档")
    print("=" * 80)
    
    doc = Document(str(paper_path))
    
    # 统计信息
    placeholders = []
    tables_count = len(doc.tables)
    paragraphs_count = len(doc.paragraphs)
    
    print(f"\n文档基本信息:")
    print(f"  - 段落数: {paragraphs_count}")
    print(f"  - 表格数: {tables_count}")
    
    # 检查段落中的占位符
    print("\n" + "=" * 80)
    print("检查段落中的占位符")
    print("=" * 80)
    
    placeholder_patterns = [
        r'图\s*\d+',  # 图 1, 图2等
        r'表\s*\d+',  # 表 1, 表2等
        r'Figure\s*\d+',
        r'Table\s*\d+',
        r'⬜',  # 方框占位符
        r'XXX',  # XXX占位符
        r'待补充',
        r'待更新',
        r'placeholder',
        r'TODO',
        r'\[.*?\]',  # 方括号内容
    ]
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
            
        for pattern in placeholder_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                placeholders.append({
                    'type': 'paragraph',
                    'index': i,
                    'text': text[:100],  # 只显示前100字符
                    'match': match.group(),
                    'position': match.start()
                })
                print(f"段落 {i}: 发现占位符 '{match.group()}'")
                print(f"  内容: {text[:100]}...")
                print()
    
    # 检查表格中的占位符
    print("\n" + "=" * 80)
    print("检查表格中的占位符")
    print("=" * 80)
    
    for table_idx, table in enumerate(doc.tables):
        print(f"\n表格 {table_idx}:")
        empty_cells = 0
        placeholder_cells = 0
        
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                
                # 检查是否为空
                if not cell_text:
                    empty_cells += 1
                    continue
                
                # 检查是否包含占位符
                for pattern in placeholder_patterns:
                    if re.search(pattern, cell_text, re.IGNORECASE):
                        placeholder_cells += 1
                        print(f"  [{row_idx},{col_idx}] 占位符: {cell_text[:50]}")
                        placeholders.append({
                            'type': 'table',
                            'table_index': table_idx,
                            'row': row_idx,
                            'col': col_idx,
                            'text': cell_text[:100],
                            'match': re.search(pattern, cell_text, re.IGNORECASE).group()
                        })
                        break
        
        print(f"  空单元格: {empty_cells}")
        print(f"  占位符单元格: {placeholder_cells}")
    
    # 检查图片
    print("\n" + "=" * 80)
    print("检查图片")
    print("=" * 80)
    
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_count += 1
    
    print(f"文档中的图片数量: {image_count}")
    
    # 总结
    print("\n" + "=" * 80)
    print("总结")
    print("=" * 80)
    print(f"总占位符数量: {len(placeholders)}")
    print(f"  - 段落占位符: {sum(1 for p in placeholders if p['type'] == 'paragraph')}")
    print(f"  - 表格占位符: {sum(1 for p in placeholders if p['type'] == 'table')}")
    print(f"  - 图片数量: {image_count}")
    
    if placeholders:
        print("\n需要处理的占位符:")
        for p in placeholders[:20]:  # 只显示前20个
            if p['type'] == 'paragraph':
                print(f"  - 段落 {p['index']}: {p['match']}")
            else:
                print(f"  - 表格 {p['table_index']} [{p['row']},{p['col']}]: {p['match']}")
        
        if len(placeholders) > 20:
            print(f"  ... 还有 {len(placeholders) - 20} 个占位符")
    
    return placeholders

if __name__ == "__main__":
    analyze_paper()
