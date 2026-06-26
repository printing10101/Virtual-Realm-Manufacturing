"""
分析论文表格结构并更新表格数据
"""

from docx import Document
from pathlib import Path
import json
import numpy as np

def format_mae(value):
    """格式化MAE值"""
    return f"{value:.4f}"

def format_r2(value):
    """格式化R²值"""
    return f"{value:.4f}"

def format_pcc(value):
    """格式化PCC值"""
    return f"{value:.4f}"

def analyze_paper_structure():
    """分析论文结构，定位表格位置"""
    
    paper_path = Path("../../docs/DL-LNN-论文-更新版.docx")
    if not paper_path.exists():
        print(f"错误：论文文件不存在: {paper_path}")
        return None
    
    print("=" * 80)
    print("分析论文结构")
    print("=" * 80)
    
    doc = Document(str(paper_path))
    
    # 统计段落和表格
    print(f"\n总段落数: {len(doc.paragraphs)}")
    print(f"总表格数: {len(doc.tables)}")
    
    # 查找表格前的段落，识别表格用途
    print("\n" + "=" * 80)
    print("表格位置分析")
    print("=" * 80)
    
    for i, table in enumerate(doc.tables):
        # 查找表格前的段落
        table_idx = doc.element.body.index(table._element)
        
        # 向前查找最近的段落
        prev_para = None
        for j in range(table_idx - 1, -1, -1):
            if doc.element.body[j].tag.endswith('p'):
                prev_para = doc.paragraphs[j]
                break
        
        if prev_para:
            print(f"\n表格 {i}:")
            print(f"  前一段落: {prev_para.text[:100]}")
            print(f"  表格行数: {len(table.rows)}")
            print(f"  表格列数: {len(table.columns)}")
            
            # 显示表格前3行内容
            print(f"  表格预览:")
            for row_idx in range(min(3, len(table.rows))):
                row_text = [cell.text[:20] for cell in table.rows[row_idx].cells]
                print(f"    行{row_idx}: {row_text}")
    
    return doc

def update_table2(doc, main_results):
    """更新表2：主对比实验结果"""
    
    print("\n" + "=" * 80)
    print("更新表2：主对比实验结果")
    print("=" * 80)
    
    # 查找表2（通常是第2个或第3个表格）
    table2 = None
    for i, table in enumerate(doc.tables):
        table_idx = doc.element.body.index(table._element)
        
        # 向前查找包含"表2"或"Table 2"的段落
        for j in range(table_idx - 1, max(0, table_idx - 5), -1):
            if doc.element.body[j].tag.endswith('p'):
                para = doc.paragraphs[j]
                if '表2' in para.text or 'Table 2' in para.text or '主对比' in para.text:
                    table2 = table
                    print(f"找到表2，索引: {i}")
                    break
        if table2:
            break
    
    if not table2:
        print("警告：未找到表2，尝试使用第2个表格")
        if len(doc.tables) >= 2:
            table2 = doc.tables[1]
        else:
            print("错误：表格数量不足")
            return False
    
    # 准备数据
    datasets = ['PHM2010', 'NUAA', 'NIST', 'Benchmark-1', '自采6061-T6']
    models = ['CT-LTC', 'LSTM', 'GRU', 'Transformer', 'CNN', 'PINN', 'gPINN', 'PeRCNN', 'BPNN']
    
    print(f"\n表格结构: {len(table2.rows)} 行 × {len(table2.columns)} 列")
    print(f"预期结构: 6 行 × 10 列 (1行表头 + 5行数据)")
    
    # 检查表格结构
    if len(table2.rows) < 6 or len(table2.columns) < 10:
        print(f"警告：表格结构不匹配，尝试继续更新")
    
    # 更新表头（如果需要）
    try:
        header_cells = table2.rows[0].cells
        if len(header_cells) >= 10:
            header_cells[0].text = 'Dataset'
            for idx, model in enumerate(models):
                header_cells[idx + 1].text = model
            print("表头已更新")
    except Exception as e:
        print(f"更新表头时出错: {e}")
    
    # 更新数据行
    for row_idx, dataset in enumerate(datasets):
        if row_idx + 1 >= len(table2.rows):
            print(f"警告：行索引 {row_idx + 1} 超出范围")
            break
        
        row = table2.rows[row_idx + 1]
        cells = row.cells
        
        if len(cells) < 10:
            print(f"警告：行 {row_idx + 1} 列数不足")
            continue
        
        # 更新数据集名称
        cells[0].text = dataset
        
        # 更新各模型的MAE值
        for model_idx, model in enumerate(models):
            if model_idx + 1 >= len(cells):
                break
            
            if dataset in main_results and model in main_results[dataset]:
                mae_value = main_results[dataset][model]['MAE']
                cells[model_idx + 1].text = format_mae(mae_value)
        
        print(f"行 {row_idx + 1} ({dataset}) 已更新")
    
    print("表2更新完成")
    return True

def update_table6(doc, time_constant_data):
    """更新表6：时间常数分析"""
    
    print("\n" + "=" * 80)
    print("更新表6：时间常数分析")
    print("=" * 80)
    
    # 查找表6
    table6 = None
    for i, table in enumerate(doc.tables):
        table_idx = doc.element.body.index(table._element)
        
        # 向前查找包含"表6"或"Table 6"的段落
        for j in range(table_idx - 1, max(0, table_idx - 5), -1):
            if doc.element.body[j].tag.endswith('p'):
                para = doc.paragraphs[j]
                if '表6' in para.text or 'Table 6' in para.text or '时间常数' in para.text:
                    table6 = table
                    print(f"找到表6，索引: {i}")
                    break
        if table6:
            break
    
    if not table6:
        print("警告：未找到表6，尝试使用最后一个表格")
        if len(doc.tables) >= 1:
            table6 = doc.tables[-1]
        else:
            print("错误：没有表格")
            return False
    
    print(f"\n表格结构: {len(table6.rows)} 行 × {len(table6.columns)} 列")
    
    # 更新表格数据
    layers = time_constant_data['layers']
    
    # 更新表头（如果需要）
    try:
        header_cells = table6.rows[0].cells
        if len(header_cells) >= 6:
            header_cells[0].text = 'Layer'
            header_cells[1].text = 'τ_mean'
            header_cells[2].text = 'τ_std'
            header_cells[3].text = 'τ_min'
            header_cells[4].text = 'τ_max'
            header_cells[5].text = 'τ_median'
            print("表头已更新")
    except Exception as e:
        print(f"更新表头时出错: {e}")
    
    # 更新数据行
    for layer_idx, layer_data in enumerate(layers):
        if layer_idx + 1 >= len(table6.rows):
            print(f"警告：行索引 {layer_idx + 1} 超出范围")
            break
        
        row = table6.rows[layer_idx + 1]
        cells = row.cells
        
        if len(cells) < 6:
            print(f"警告：行 {layer_idx + 1} 列数不足")
            continue
        
        # 更新数据
        cells[0].text = str(layer_data['layer'])
        cells[1].text = format_mae(layer_data['tau_mean'])
        cells[2].text = format_mae(layer_data['tau_std'])
        cells[3].text = format_mae(layer_data['tau_min'])
        cells[4].text = format_mae(layer_data['tau_max'])
        cells[5].text = format_mae(layer_data['tau_median'])
        
        print(f"层 {layer_data['layer']} 数据已更新")
    
    # 添加全局统计（如果有额外行）
    if len(table6.rows) > len(layers) + 1:
        global_row = table6.rows[len(layers) + 1]
        global_cells = global_row.cells
        
        if len(global_cells) >= 6:
            global_data = time_constant_data['global']
            global_cells[0].text = 'Global'
            global_cells[1].text = format_mae(global_data['tau_mean'])
            global_cells[2].text = format_mae(global_data['tau_std'])
            global_cells[3].text = format_mae(global_data['tau_min'])
            global_cells[4].text = format_mae(global_data['tau_max'])
            global_cells[5].text = format_mae(global_data['tau_median'])
            print("全局统计数据已更新")
    
    print("表6更新完成")
    return True

def main():
    """主函数"""
    
    print("=" * 80)
    print("论文表格更新工具")
    print("=" * 80)
    
    # 加载实验结果
    results_dir = Path("results")
    
    main_results_path = results_dir / "main_comparison_results.json"
    time_constant_path = results_dir / "time_constant_analysis.json"
    
    if not main_results_path.exists():
        print(f"错误：主对比结果文件不存在: {main_results_path}")
        return
    
    if not time_constant_path.exists():
        print(f"错误：时间常数分析文件不存在: {time_constant_path}")
        return
    
    print("\n加载实验结果...")
    with open(main_results_path, 'r', encoding='utf-8') as f:
        main_results = json.load(f)
    
    with open(time_constant_path, 'r', encoding='utf-8') as f:
        time_constant_data = json.load(f)
    
    print("✓ 主对比结果已加载")
    print("✓ 时间常数分析已加载")
    
    # 分析论文结构
    doc = analyze_paper_structure()
    if not doc:
        return
    
    # 更新表2
    update_table2(doc, main_results)
    
    # 更新表6
    update_table6(doc, time_constant_data)
    
    # 保存更新后的论文
    output_path = Path("../../docs/DL-LNN-论文-最终版.docx")
    doc.save(str(output_path))
    
    print("\n" + "=" * 80)
    print("更新完成！")
    print("=" * 80)
    print(f"\n更新后的论文已保存到: {output_path}")
    print("\n请手动检查表格内容是否正确")

if __name__ == "__main__":
    main()
