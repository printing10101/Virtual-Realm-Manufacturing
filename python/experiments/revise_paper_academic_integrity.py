"""修订论文 docx 中的学术诚信问题。

学术诚信修订说明（重要）：
    本脚本用于修订两篇论文 docx 中与代码实现不一致的声称，
    确保"论文声称"与"代码实现"保持一致。修订原则：
      1. PHM2010 数据集：代码已接入真实公开数据（UniwearDataLoader），
         论文需相应更新，明确标注 PHM2010 为真实公开数据集。
      2. 6061-T6 数据集：代码中 Industrial6061T6Dataset 为合成数据占位实现，
         论文中所有"自采 6061-T6 工业数据集"措辞必须修改为
         "合成 6061-T6 占位数据集"，不可声称对应真实自采数据。
      3. 移除"实际颤振发生的临界切深 1.38 mm"这种暗示真实实验的措辞，
         改为"解析模型计算的临界切深"。

修订范围：
    - docs/大创赛/论文实验/DL-LNN论文_数据修正版_仿真数据说明版.docx
    - docs/大创赛/论文实验/DL-LNN综合实验报告.docx

备份策略：
    修改前自动备份原始文件为 *.docx.bak
"""
from __future__ import annotations

import logging
import shutil
from pathlib import Path
from typing import List, Tuple

from docx import Document

logger = logging.getLogger(__name__)

_PROJECT_ROOT = Path(__file__).resolve().parents[2]
_PAPERS_DIR = _PROJECT_ROOT / "docs" / "大创赛" / "论文实验"

_PAPER_FILES = [
    _PAPERS_DIR / "DL-LNN论文_数据修正版_仿真数据说明版.docx",
    _PAPERS_DIR / "DL-LNN综合实验报告.docx",
]

# 通用的文本替换规则（应用于两篇论文的所有段落）
# 格式：(原文片段, 替换后片段, 说明)
_COMMON_REPLACEMENTS: List[Tuple[str, str, str]] = [
    # —— 6061-T6 措辞修正：将"自采"修改为"合成占位" ——
    (
        "自采6061-T6工业数据集",
        "合成6061-T6占位数据集",
        "6061-T6 不可声称自采",
    ),
    (
        "自采 6061-T6 工业数据集",
        "合成 6061-T6 占位数据集",
        "6061-T6 不可声称自采",
    ),
    (
        "自采6061-T6铝合金工业数据集",
        "合成6061-T6铝合金占位数据集",
        "6061-T6 不可声称自采",
    ),
    (
        "自采 6061-T6 铝合金数据集",
        "合成 6061-T6 铝合金占位数据集",
        "6061-T6 不可声称自采",
    ),
    (
        "自采6061-T6数据集",
        "合成6061-T6占位数据集",
        "6061-T6 不可声称自采",
    ),
    (
        "自采 6061-T6 数据集",
        "合成 6061-T6 占位数据集",
        "6061-T6 不可声称自采",
    ),
    (
        "6061-T6 自采",
        "6061-T6 合成占位",
        "6061-T6 不可声称自采",
    ),
    # —— "实际颤振发生的临界切深"修改为"解析模型计算的临界切深" ——
    (
        "实际颤振发生的临界切深",
        "解析模型计算的临界切深",
        "移除暗示真实实验的措辞",
    ),
    # —— "本项目自采的6061-T6"修改为"合成占位" ——
    (
        "本项目自采的6061-T6铝合金工业数据集",
        "合成6061-T6铝合金占位数据集（基于TlustyAnalyticalModel生成，非真实自采数据）",
        "明确标注为合成占位",
    ),
    (
        "本项目自采的 6061-T6 铝合金工业数据集",
        "合成 6061-T6 铝合金占位数据集（基于 TlustyAnalyticalModel 生成，非真实自采数据）",
        "明确标注为合成占位",
    ),
]


def _backup_file(file_path: Path) -> Path:
    """备份原始文件为 *.docx.bak。"""
    backup_path = file_path.with_suffix(".docx.bak")
    if not backup_path.exists():
        shutil.copy2(file_path, backup_path)
        logger.info("已备份原始文件: %s -> %s", file_path.name, backup_path.name)
    else:
        logger.info("备份文件已存在，跳过备份: %s", backup_path.name)
    return backup_path


def _apply_replacements_to_paragraph(paragraph, replacements: List[Tuple[str, str, str]]) -> int:
    """对单个段落应用文本替换，返回替换次数。

    采用"段落级合并替换"策略：
        1. 将段落所有 run 的文本拼接为完整文本
        2. 检查是否包含需要替换的片段
        3. 若包含，将第一个 run 的文本设为替换后的完整文本，其余 run 清空
        4. 这样会丢失段内 run 级格式，但保留段落级格式（缩进、对齐等）

    Args:
        paragraph: python-docx 的 Paragraph 对象
        replacements: 替换规则列表

    Returns:
        本段落应用的替换次数
    """
    if not paragraph.runs:
        return 0

    full_text = "".join(run.text for run in paragraph.runs)
    if not full_text.strip():
        return 0

    new_text = full_text
    count = 0
    for old, new, _desc in replacements:
        if old in new_text:
            new_text = new_text.replace(old, new)
            count += 1

    if count > 0 and new_text != full_text:
        # 将替换后的文本写入第一个 run，其余 run 清空
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""

    return count


def _revise_paper(file_path: Path) -> dict:
    """修订单篇论文 docx。

    Args:
        file_path: 论文 docx 文件路径

    Returns:
        修订统计信息字典
    """
    logger.info("开始修订论文: %s", file_path.name)

    # 备份
    _backup_file(file_path)

    doc = Document(str(file_path))

    total_paragraphs = 0
    revised_paragraphs = 0
    total_replacements = 0
    revised_details: list[dict] = []

    for idx, paragraph in enumerate(doc.paragraphs):
        total_paragraphs += 1
        text_before = paragraph.text.strip()

        if not text_before:
            continue

        count = _apply_replacements_to_paragraph(paragraph, _COMMON_REPLACEMENTS)

        if count > 0:
            revised_paragraphs += 1
            total_replacements += count
            text_after = paragraph.text.strip()
            revised_details.append({
                "paragraph_index": idx,
                "replacements_applied": count,
                "text_before": text_before[:200],
                "text_after": text_after[:200],
            })

    # 保存修订后的文件
    doc.save(str(file_path))

    stats = {
        "file": file_path.name,
        "total_paragraphs": total_paragraphs,
        "revised_paragraphs": revised_paragraphs,
        "total_replacements": total_replacements,
        "details": revised_details,
    }

    logger.info(
        "修订完成: %s（%d/%d 段落修订，共 %d 处替换）",
        file_path.name,
        revised_paragraphs,
        total_paragraphs,
        total_replacements,
    )

    return stats


def _apply_special_revisions_paper1(file_path: Path) -> dict:
    """对第一篇论文应用特殊修订（段落级完全重写）。

    这些修订涉及需要完全重写的段落，无法通过简单文本替换完成。

    Args:
        file_path: 论文 docx 文件路径

    Returns:
        修订统计信息字典
    """
    logger.info("应用第一篇论文的特殊修订: %s", file_path.name)

    doc = Document(str(file_path))

    special_revisions = [
        # (匹配片段, 替换后的完整段落文本, 说明)
        (
            "本文使用的五个数据集均为仿真生成数据",
            "本文使用的数据集中，PHM2010 为真实公开数据集（通过 UniwearDataLoader 加载真实铣削信号，3 组实验 c1/c4/c6，不锈钢 HRC52，共 104675 行），其余 4 个数据集（NUAA、NIST、Benchmark-1、6061-T6）为基于 Tlusty 解析模型生成的仿真数据。具体参数设置如下：",
            "PHM2010 为真实数据，不可声称全部仿真",
        ),
        (
            "本文实验所使用的数据集（PHM2010、NUAA、NIST、Benchmark-1、6061-T6）均为基于Tlusty再生颤振理论生成的仿真数据",
            "本文实验所使用的数据集中，PHM2010 为真实公开数据集（PHM Society 2010 切削颤振挑战赛），NUAA、NIST、Benchmark-1 为基于 Tlusty 再生颤振理论生成的仿真数据，6061-T6 为合成占位数据集（基于 TlustyAnalyticalModel 生成，非真实自采数据）",
            "PHM2010 为真实数据，6061-T6 为合成占位",
        ),
    ]

    revised_count = 0
    for idx, paragraph in enumerate(doc.paragraphs):
        full_text = paragraph.text.strip()
        if not full_text:
            continue

        for match_text, new_text, desc in special_revisions:
            if match_text in full_text:
                # 完全重写段落
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                    for run in paragraph.runs[1:]:
                        run.text = ""
                    revised_count += 1
                    logger.info(
                        "特殊修订 [%d]: %s\n  原文: %s...\n  新文: %s...",
                        idx,
                        desc,
                        full_text[:100],
                        new_text[:100],
                    )

    doc.save(str(file_path))
    return {"file": file_path.name, "special_revisions": revised_count}


def main() -> None:
    """主入口：修订两篇论文 docx 中的学术诚信问题。"""
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(message)s",
    )

    print("=" * 80)
    print("论文 docx 学术诚信修订")
    print("=" * 80)
    print()
    print("修订原则：")
    print("  1. PHM2010 数据集：代码已接入真实公开数据，论文需明确标注为真实数据")
    print("  2. 6061-T6 数据集：代码为合成占位实现，论文需修改'自采'措辞为'合成占位'")
    print("  3. 移除暗示真实实验的措辞（如'实际颤振发生的临界切深'）")
    print()

    all_stats = []

    for paper_file in _PAPER_FILES:
        if not paper_file.exists():
            logger.warning("论文文件不存在，跳过: %s", paper_file)
            continue

        # 通用文本替换
        stats = _revise_paper(paper_file)
        all_stats.append(stats)

        # 第一篇论文的特殊修订（段落级完全重写）
        if "DL-LNN论文" in paper_file.name:
            special_stats = _apply_special_revisions_paper1(paper_file)
            all_stats.append(special_stats)

    print()
    print("=" * 80)
    print("修订汇总")
    print("=" * 80)
    for stats in all_stats:
        print(f"\n文件: {stats.get('file', 'N/A')}")
        if "total_replacements" in stats:
            print(f"  总段落数: {stats['total_paragraphs']}")
            print(f"  修订段落数: {stats['revised_paragraphs']}")
            print(f"  总替换次数: {stats['total_replacements']}")
            if stats["details"]:
                print("  修订详情（前 10 条）:")
                for detail in stats["details"][:10]:
                    print(f"    [段落 {detail['paragraph_index']}] {detail['replacements_applied']} 处替换")
                    print(f"      原文: {detail['text_before'][:120]}...")
                    print(f"      新文: {detail['text_after'][:120]}...")
        if "special_revisions" in stats:
            print(f"  特殊修订次数: {stats['special_revisions']}")

    print()
    print("=" * 80)
    print("修订完成！原始文件已备份为 *.docx.bak")
    print("=" * 80)


if __name__ == "__main__":
    main()
