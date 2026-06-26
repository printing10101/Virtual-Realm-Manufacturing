#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
从论文DOCX文件中提取所有实验、图表、数据集、模型、指标信息
"""

from docx import Document
import re
import json

def extract_paper_info(docx_path):
    """提取论文信息"""
    doc = Document(docx_path)
    
    # 提取所有文本
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())
    
    # 提取表格
    tables_text = []
    for i, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        tables_text.append({
            'table_index': i,
            'data': table_data
        })
    
    # 合并所有文本用于搜索
    all_text = '\n'.join(full_text)
    
    # 提取实验名称和编号
    experiments = []
    exp_patterns = [
        r'实验\s*(\d+)[：:]\s*([^\n]+)',
        r'Experiment\s*(\d+)[：:]\s*([^\n]+)',
        r'(\d+\.\d+)\s+([^\n]*实验[^\n]*)',
    ]
    for pattern in exp_patterns:
        matches = re.finditer(pattern, all_text)
        for match in matches:
            experiments.append(match.group(0))
    
    # 提取图表编号
    figures = re.findall(r'图\s*(\d+)', all_text)
    tables = re.findall(r'表\s*(\d+)', all_text)
    
    # 提取数据集
    datasets = []
    dataset_patterns = [
        r'PHM2010',
        r'NUAA',
        r'NIST',
        r'Benchmark-\d+',
        r'6061-T6',
        r'7075-T6',
        r'2024-T3',
        r'304SS',
        r'Ti6Al4V',
        r'合成数据集',
        r'工业数据集',
    ]
    for pattern in dataset_patterns:
        if re.search(pattern, all_text, re.IGNORECASE):
            datasets.append(pattern)
    
    # 提取模型/方法
    models = []
    model_patterns = [
        r'CT-LTC',
        r'LSTM',
        r'GRU',
        r'Transformer',
        r'CNN',
        r'PINN',
        r'gPINN',
        r'PeRCNN',
        r'BPNN',
    ]
    for pattern in model_patterns:
        if re.search(pattern, all_text):
            models.append(pattern)
    
    # 提取评估指标
    metrics = []
    metric_patterns = [
        r'MAE',
        r'RMSE',
        r'R²',
        r'R2',
        r'MAPE',
        r'PCC',
    ]
    for pattern in metric_patterns:
        if re.search(pattern, all_text):
            metrics.append(pattern)
    
    # 查找占位符
    placeholders = []
    placeholder_patterns = [
        r'⬜',
        r'TBD',
        r'待填写',
        r'待补充',
        r'XXX',
        r'\[.*?\]',  # 方括号占位符
    ]
    for pattern in placeholder_patterns:
        matches = re.findall(pattern, all_text)
        if matches:
            placeholders.extend(matches)
    
    # 提取关键数值
    values = []
    value_pattern = r'(\d+\.\d+)'
    value_matches = re.findall(value_pattern, all_text)
    
    return {
        'experiments': list(set(experiments))[:20],  # 限制数量
        'figures': sorted(list(set(figures)), key=int),
        'tables': sorted(list(set(tables)), key=int),
        'datasets': list(set(datasets)),
        'models': list(set(models)),
        'metrics': list(set(metrics)),
        'placeholders': placeholders[:10],  # 限制数量
        'sample_values': value_matches[:20],  # 样本数值
        'total_paragraphs': len(full_text),
        'total_tables': len(doc.tables)
    }

if __name__ == '__main__':
    docx_path = 'docs/DL-LNN-论文-最终版.docx'
    info = extract_paper_info(docx_path)
    
    print("=" * 60)
    print("论文信息提取结果")
    print("=" * 60)
    print(f"\n总段落数: {info['total_paragraphs']}")
    print(f"总表格数: {info['total_tables']}")
    
    print(f"\n实验 ({len(info['experiments'])}个):")
    for exp in info['experiments']:
        print(f"  - {exp}")
    
    print(f"\n图表编号:")
    print(f"  图: {', '.join(info['figures'])}")
    print(f"  表: {', '.join(info['tables'])}")
    
    print(f"\n数据集 ({len(info['datasets'])}个):")
    for ds in info['datasets']:
        print(f"  - {ds}")
    
    print(f"\n模型/方法 ({len(info['models'])}个):")
    for model in info['models']:
        print(f"  - {model}")
    
    print(f"\n评估指标 ({len(info['metrics'])}个):")
    for metric in info['metrics']:
        print(f"  - {metric}")
    
    print(f"\n占位符 ({len(info['placeholders'])}个):")
    for ph in info['placeholders']:
        print(f"  - {ph}")
    
    # 保存为JSON
    with open('paper_info.json', 'w', encoding='utf-8') as f:
        json.dump(info, f, ensure_ascii=False, indent=2)
    
    print(f"\n详细信息已保存到: paper_info.json")
