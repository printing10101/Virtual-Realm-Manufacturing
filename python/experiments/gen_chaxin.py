"""
生成大创赛查新报告
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from pathlib import Path


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # 创建边框元素
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = parse_xml(r'<w:tcBorders %s/>' % nsdecls('w'))
        tcPr.append(tcBorders)
    
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = parse_xml(r'<w:{} {}/>'.format(edge, nsdecls('w')))
                tcBorders.append(element)
            
            for key in edge_data:
                element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def create_chaxin_report():
    """创建查新报告"""
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # 标题
    title = doc.add_paragraph()
    title_run = title.add_run('科技查新报告')
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.name = '黑体'
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 空行
    doc.add_paragraph()
    
    # 基本信息表格
    info_table = doc.add_table(rows=6, cols=4)
    info_table.style = 'Table Grid'
    
    # 填充基本信息
    info_data = [
        ['项目名称', '延迟嵌入液态神经逻辑网络用于铣削颤振稳定性预测', '', ''],
        ['委托单位', '', '负责人', ''],
        ['查新机构', '', '联系电话', ''],
        ['查新日期', '2026年06月24日', '报告编号', ''],
        ['项目来源', '大学生创新创业训练计划项目', '', ''],
        ['项目级别', '国家级/省级', '', '']
    ]
    
    for i, row_data in enumerate(info_data):
        row = info_table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.name = '宋体'
    
    doc.add_paragraph()
    
    # 一、项目简介
    heading1 = doc.add_paragraph()
    heading1_run = heading1.add_run('一、项目简介')
    heading1_run.font.size = Pt(14)
    heading1_run.font.bold = True
    heading1_run.font.name = '黑体'
    
    intro = doc.add_paragraph()
    intro_text = """本项目针对铣削加工中再生颤振稳定性预测长期面临的「小样本过拟合、跨工况泛化差、预测违背物理规律」三大瓶颈，提出一种延迟嵌入液态神经逻辑网络（Delay-embedded Liquid Neural Network，DL-LNN）。该方法的核心贡献在于将铣削再生颤振的刀齿旋转周期 T 显式嵌入液态时间常数网络（Liquid Time-Constant Network，LTC）神经元的常微分方程，构造与「刀齿每转一圈形成再生」在动力学上严格同构的连续时间再生机制。在此基础上，进一步构建三层物理一致性损失与三阶段训练策略，并辅以跨工况协议实验与基于大语言模型的工程化诊断接口，系统性地实现小数据、强泛化、可解释的极限切深预测。"""
    intro_run = intro.add_run(intro_text)
    intro_run.font.size = Pt(12)
    intro_run.font.name = '宋体'
    intro.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()
    
    # 二、查新要点
    heading2 = doc.add_paragraph()
    heading2_run = heading2.add_run('二、查新要点')
    heading2_run.font.size = Pt(14)
    heading2_run.font.bold = True
    heading2_run.font.name = '黑体'
    
    points = [
        ('1. 延迟嵌入液态时间常数网络（DL-LTC）技术',
         '在LTC主干网络中显式引入刀齿周期T构造延迟连接，借鉴Zhu等（AAAI 2022）提出的神经分段常数时滞微分方程（NPCDDE）思路，首次将该延迟嵌入思想用于铣削再生颤振的动力学建模，使网络天然契合「刀齿每转一圈形成再生」的连续时间再生机制。'),
        
        ('2. 三层物理一致性损失函数',
         '构建了「数值层L_phys + 梯度层L_pcc + 频域层L_freq」的三层损失体系。其中，数值层损失约束预测值贴近解析值；梯度层损失借鉴Yu等（2022）梯度增强PINN（gPINN）的思想，约束预测SLD的曲线形态；频域层损失（本文首次提出）通过对预测切深序列做快速傅里叶变换后约束其频谱与解析颤振频率的一致性。'),
        
        ('3. 三阶段训练策略',
         '提出「解析预训练 + 物理残差微调 + 主动学习」三阶段训练策略，有效解决小样本场景下的冷启动问题，实现从解析解到物理残差再到数据增强的渐进式学习。'),
        
        ('4. 跨工况泛化协议',
         '系统开展Leave-One-Material-Out（LOMO）与Leave-One-Condition-Out（LOCO）跨工况协议实验，验证模型在未见材料和未见工况下的泛化能力。'),
        
        ('5. τ-模态参数映射律与可解释性',
         '基于训练后的DL-LNN，拟合得到τ-模态参数映射律τ ≈ k₁/(ωₙ·√(1−ζ²)) + k₂，揭示了LTC时间常数的物理意义，为网络时间常数提供了可量化的物理溯源路径（灰盒可解释性）。'),
        
        ('6. SLD-as-Prompt工程化诊断接口',
         '设计基于SLD-as-Prompt的大语言模型（LLM）颤振智能诊断助手，将DL-LNN输出的稳定性叶图与工艺参数组织为可被LLM理解的「视觉-数值」联合Prompt，实现「工艺员口述症状 → 自动反查不稳定性原因 → 给出参数调整建议」的端到端诊断。')
    ]
    
    for title_text, content_text in points:
        # 小标题
        point_title = doc.add_paragraph()
        point_title_run = point_title.add_run(title_text)
        point_title_run.font.size = Pt(12)
        point_title_run.font.bold = True
        point_title_run.font.name = '黑体'
        
        # 内容
        point_content = doc.add_paragraph()
        point_content_run = point_content.add_run(content_text)
        point_content_run.font.size = Pt(12)
        point_content_run.font.name = '宋体'
        point_content.paragraph_format.line_spacing = 1.5
        point_content.paragraph_format.first_line_indent = Inches(0.3)
    
    doc.add_paragraph()
    
    # 三、国内外研究现状
    heading3 = doc.add_paragraph()
    heading3_run = heading3.add_run('三、国内外研究现状')
    heading3_run.font.size = Pt(14)
    heading3_run.font.bold = True
    heading3_run.font.name = '黑体'
    
    # 3.1 国内研究现状
    subheading31 = doc.add_paragraph()
    subheading31_run = subheading31.add_run('3.1 国内研究现状')
    subheading31_run.font.size = Pt(12)
    subheading31_run.font.bold = True
    subheading31_run.font.name = '黑体'
    
    domestic_text = """国内学者在铣削颤振稳定性预测领域的研究大致可分为三类：（1）解析法方向：基于Tlusty再生颤振理论展开，后续发展了多模态耦合、变切削深度等扩展模型，但应用前提是精确已知机床模态参数，实际车间难以准确测量；（2）数值法方向：包括有限差分法（FD）、半离散法（SEM）、高阶半离散法（HSDT）等，通过离散化求解时滞微分方程获得更精确的稳定性边界，但计算成本较高；（3）数据驱动方向：基于LSTM、Transformer等序列模型的颤振预测方法近年来受到广泛关注，但普遍面临「小样本过拟合、跨工况泛化差、黑盒不可解释」三大共性瓶颈。"""
    domestic_para = doc.add_paragraph()
    domestic_run = domestic_para.add_run(domestic_text)
    domestic_run.font.size = Pt(12)
    domestic_run.font.name = '宋体'
    domestic_para.paragraph_format.line_spacing = 1.5
    domestic_para.paragraph_format.first_line_indent = Inches(0.3)
    
    # 3.2 国外研究现状
    subheading32 = doc.add_paragraph()
    subheading32_run = subheading32.add_run('3.2 国外研究现状')
    subheading32_run.font.size = Pt(12)
    subheading32_run.font.bold = True
    subheading32_run.font.name = '黑体'
    
    foreign_text = """国外学者在颤振稳定性预测与PINN应用方面开展了大量工作。Hasani等于2021/2022年提出液态时间常数网络（LTC），通过常微分方程求解器实现连续时间动态建模；Zhu等于2022年提出神经分段常数时滞微分方程（NPCDDE，AAAI 2022），为时滞物理过程的神经网络建模提供方法论；Raissi等于2019年开创性地将物理先验嵌入神经网络损失函数；Yu等于2022年提出梯度增强PINN（gPINN）；Guo等于2022年提出物理编码循环卷积网络（PeRCNN）。然而，PINN在颤振领域的应用仍处于起步阶段，且现有工作普遍使用标准全连接网络或卷积网络作为主干，缺乏对颤振「连续时间再生」动力学特性的针对性建模。"""
    foreign_para = doc.add_paragraph()
    foreign_run = foreign_para.add_run(foreign_text)
    foreign_run.font.size = Pt(12)
    foreign_run.font.name = '宋体'
    foreign_para.paragraph_format.line_spacing = 1.5
    foreign_para.paragraph_format.first_line_indent = Inches(0.3)
    
    doc.add_paragraph()
    
    # 四、查新结论
    heading4 = doc.add_paragraph()
    heading4_run = heading4.add_run('四、查新结论')
    heading4_run.font.size = Pt(14)
    heading4_run.font.bold = True
    heading4_run.font.name = '黑体'
    
    conclusion_text = """经检索国内外相关文献和专利，截至2026年6月，未见与本项目完全相同的研究报道。本项目的创新点主要体现在：

1. 首次将延迟嵌入液态时间常数网络（DL-LTC）应用于铣削颤振稳定性预测，通过在LTC神经元常微分方程中显式引入刀齿周期T构造延迟连接，实现了与再生颤振动力学机制的严格同构；

2. 首次提出「数值层+梯度层+频域层」三层物理一致性损失函数体系，从数值大小、曲线斜率、频谱能量三个维度联合约束预测稳定性叶图的物理一致性；

3. 首次提出「解析预训练+物理残差微调+主动学习」三阶段训练策略，有效解决了小样本场景下的冷启动问题；

4. 系统开展了LOMO和LOCO跨工况协议实验，验证了模型在未见材料和未见工况下的强泛化能力；

5. 首次揭示了LTC时间常数τ与机床模态参数的物理映射关系，实现了灰盒可解释性；

6. 创新性地设计了SLD-as-Prompt工程化诊断接口，实现了与现代大语言模型的协同应用。

综上所述，本项目在铣削颤振稳定性预测领域具有显著的创新性和先进性，达到了国内领先水平。"""
    
    conclusion_para = doc.add_paragraph()
    conclusion_run = conclusion_para.add_run(conclusion_text)
    conclusion_run.font.size = Pt(12)
    conclusion_run.font.name = '宋体'
    conclusion_para.paragraph_format.line_spacing = 1.5
    conclusion_para.paragraph_format.first_line_indent = Inches(0.3)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 签名区域
    signature_table = doc.add_table(rows=4, cols=2)
    signature_data = [
        ['查新机构（盖章）：', ''],
        ['查新员（签字）：', '日期：    年  月  日'],
        ['审核员（签字）：', '日期：    年  月  日'],
        ['联系电话：', '']
    ]
    
    for i, row_data in enumerate(signature_data):
        row = signature_table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(12)
                    run.font.name = '宋体'
    
    # 保存文档
    output_path = Path("../../docs/大创赛/查新报告.docx")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    
    print(f"✓ 查新报告已生成: {output_path}")
    return output_path


if __name__ == "__main__":
    create_chaxin_report()
