"""
将实验图片和数据插入到论文最终版
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def insert_figures_and_data():
    """插入图片和数据到论文"""
    
    # 读取论文
    paper_path = Path("../../docs/DL-LNN-论文-最终版.docx")
    if not paper_path.exists():
        print(f"错误: 论文文件不存在: {paper_path}")
        return
    
    print("=" * 80)
    print("开始插入图片和数据到论文")
    print("=" * 80)
    
    doc = Document(str(paper_path))
    
    # 图片映射
    figure_mapping = {
        "图 2": "results/figures/main_results_synthetic.png",
        "图 3": "results/figures/active_learning_curves.png",
    }
    
    # 查找并替换图片占位符
    figures_inserted = 0
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        
        # 查找图片占位符 [图 X：...]
        match = re.search(r'\[图\s*(\d+)[：:](.*?)\]', text)
        if match:
            fig_num = match.group(1)
            fig_key = f"图 {fig_num}"
            
            if fig_key in figure_mapping:
                img_path = Path(figure_mapping[fig_key])
                if img_path.exists():
                    print(f"段落 {i}: 找到图片占位符 {fig_key}")
                    print(f"  图片路径: {img_path}")
                    
                    # 清空当前段落
                    para.clear()
                    
                    # 插入图片
                    run = para.add_run()
                    run.add_picture(str(img_path), width=Inches(5.5))
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    figures_inserted += 1
                    print(f"  ✓ 已插入图片")
                else:
                    print(f"段落 {i}: 图片文件不存在: {img_path}")
    
    print(f"\n共插入 {figures_inserted} 张图片")
    
    # 保存更新后的论文
    output_path = Path("../../docs/DL-LNN-论文-最终版-完整.docx")
    doc.save(str(output_path))
    print(f"\n论文已保存到: {output_path}")
    print("=" * 80)


if __name__ == "__main__":
    insert_figures_and_data()
