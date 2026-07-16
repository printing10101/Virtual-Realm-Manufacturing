"""
更新论文表格数据
使用实验结果更新论文中的表2、表6和表7
"""

import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def load_results():
    """加载所有实验结果"""
    results_dir = Path("results")
    
    with open(results_dir / "main_comparison_results.json", 'r', encoding='utf-8') as f:
        main_results = json.load(f)
    
    with open(results_dir / "time_constant_analysis.json", 'r', encoding='utf-8') as f:
        time_constant = json.load(f)
    
    with open(results_dir / "ablation_results.json", 'r', encoding='utf-8') as f:
        ablation = json.load(f)
    
    with open(results_dir / "cross_condition_results.json", 'r', encoding='utf-8') as f:
        cross_condition = json.load(f)
    
    return main_results, time_constant, ablation, cross_condition


def format_table2_data(main_results):
    """格式化表2数据"""
    print("=" * 80)
    print("表2: 主对比实验结果 (5个数据集 × 9个模型)")
    print("=" * 80)
    
    # MAE表格
    print("\nMAE (Mean Absolute Error):")
    print("-" * 130)
    header = f"{'Dataset':<15}" + "".join([f"{name:<13}" for name in ['DL-LNN', 'LSTM', 'GRU', 'Transformer', 'CNN', 'PINN', 'gPINN', 'PeRCNN', 'BPNN']])
    print(header)
    print("-" * 130)
    
    for dataset_name in ['PHM2010', 'NUAA', 'NIST', 'Benchmark-1', '自采6061-T6']:
        row = f"{dataset_name:<15}"
        for model_name in ['DL-LNN', 'LSTM', 'GRU', 'Transformer', 'CNN', 'PINN', 'gPINN', 'PeRCNN', 'BPNN']:
            if dataset_name in main_results and model_name in main_results[dataset_name]:
                mae = main_results[dataset_name][model_name]['MAE']
                row += f"{mae:<13.4f}"
            else:
                row += f"{'N/A':<13}"
        print(row)
    print("-" * 130)
    
    # R²表格
    print("\nR² (Coefficient of Determination):")
    print("-" * 130)
    print(header)
    print("-" * 130)
    
    for dataset_name in ['PHM2010', 'NUAA', 'NIST', 'Benchmark-1', '自采6061-T6']:
        row = f"{dataset_name:<15}"
        for model_name in ['DL-LNN', 'LSTM', 'GRU', 'Transformer', 'CNN', 'PINN', 'gPINN', 'PeRCNN', 'BPNN']:
            if dataset_name in main_results and model_name in main_results[dataset_name]:
                r2 = main_results[dataset_name][model_name]['R2']
                row += f"{r2:<13.4f}"
            else:
                row += f"{'N/A':<13}"
        print(row)
    print("-" * 130)
    
    # PCC表格
    print("\nPCC (Pearson Correlation Coefficient):")
    print("-" * 130)
    print(header)
    print("-" * 130)
    
    for dataset_name in ['PHM2010', 'NUAA', 'NIST', 'Benchmark-1', '自采6061-T6']:
        row = f"{dataset_name:<15}"
        for model_name in ['DL-LNN', 'LSTM', 'GRU', 'Transformer', 'CNN', 'PINN', 'gPINN', 'PeRCNN', 'BPNN']:
            if dataset_name in main_results and model_name in main_results[dataset_name]:
                pcc = main_results[dataset_name][model_name]['PCC']
                row += f"{pcc:<13.4f}"
            else:
                row += f"{'N/A':<13}"
        print(row)
    print("-" * 130)


def format_table6_data(time_constant):
    """格式化表6数据"""
    print("\n" + "=" * 80)
    print("表6: 时间常数分析结果")
    print("=" * 80)
    
    print("\n各层时间常数统计:")
    print("-" * 100)
    print(f"{'Layer':<10}{'τ_mean':<15}{'τ_std':<15}{'τ_min':<15}{'τ_max':<15}{'τ_median':<15}")
    print("-" * 100)
    
    for layer in time_constant['layers']:
        print(f"Layer {layer['layer']:<5}"
              f"{layer['tau_mean']:<15.4f}"
              f"{layer['tau_std']:<15.4f}"
              f"{layer['tau_min']:<15.4f}"
              f"{layer['tau_max']:<15.4f}"
              f"{layer['tau_median']:<15.4f}")
    
    print("-" * 100)
    print(f"{'Global':<10}"
          f"{time_constant['global']['tau_mean']:<15.4f}"
          f"{time_constant['global']['tau_std']:<15.4f}"
          f"{time_constant['global']['tau_min']:<15.4f}"
          f"{time_constant['global']['tau_max']:<15.4f}"
          f"{time_constant['global']['tau_median']:<15.4f}")
    print("-" * 100)
    
    # 分析时间常数分布
    all_taus = []
    for layer in time_constant['layers']:
        all_taus.extend(layer['tau_values'])
    
    fast_count = sum(1 for t in all_taus if t < 0.05)
    medium_count = sum(1 for t in all_taus if 0.05 <= t < 0.15)
    slow_count = sum(1 for t in all_taus if t >= 0.15)
    
    print(f"\n时间常数分布分析:")
    print(f"  - 快速响应单元 (τ < 0.05): {fast_count} 个 ({fast_count/len(all_taus)*100:.2f}%)")
    print(f"  - 中速响应单元 (0.05 ≤ τ < 0.15): {medium_count} 个 ({medium_count/len(all_taus)*100:.2f}%)")
    print(f"  - 慢速响应单元 (τ ≥ 0.15): {slow_count} 个 ({slow_count/len(all_taus)*100:.2f}%)")


def update_paper_tables():
    """更新论文表格"""
    print("加载实验结果...")
    main_results, time_constant, ablation, cross_condition = load_results()
    
    # 格式化表2数据
    format_table2_data(main_results)
    
    # 格式化表6数据
    format_table6_data(time_constant)
    
    print("\n" + "=" * 80)
    print("表格数据格式化完成！")
    print("=" * 80)
    
    # 生成Word表格
    print("\n生成Word表格...")
    
    # 读取论文
    paper_path = Path("../../docs/DL-LNN-论文-更新版.docx")
    if not paper_path.exists():
        print(f"警告: 论文文件不存在: {paper_path}")
        return
    
    doc = Document(str(paper_path))
    
    # 查找并更新表格
    table_count = 0
    for i, table in enumerate(doc.tables):
        # 检查是否是表2（主对比结果）
        if i > 0 and i < len(doc.tables) - 1:
            prev_para = doc.paragraphs[max(0, i - 1)]
            if '表2' in prev_para.text or 'Table 2' in prev_para.text:
                print(f"找到表2 (索引 {i})")
                update_table2(table, main_results)
                table_count += 1
            
            # 检查是否是表6（时间常数分析）
            if '表6' in prev_para.text or 'Table 6' in prev_para.text:
                print(f"找到表6 (索引 {i})")
                update_table6(table, time_constant)
                table_count += 1
    
    # 保存更新后的论文
    output_path = Path("../../docs/DL-LNN-论文-最终版.docx")
    doc.save(str(output_path))
    print(f"\n论文已保存到: {output_path}")
    print(f"更新了 {table_count} 个表格")


def update_table2(table, main_results):
    """更新表2"""
    print("更新表2...")
    
    # 清空表格内容（保留表头）
    for i in range(1, len(table.rows)):
        for j in range(len(table.rows[i].cells)):
            table.rows[i].cells[j].text = ""
    
    # 填充数据
    datasets = ['PHM2010', 'NUAA', 'NIST', 'Benchmark-1', '自采6061-T6']
    models = ['DL-LNN', 'LSTM', 'GRU', 'Transformer', 'CNN', 'PINN', 'gPINN', 'PeRCNN', 'BPNN']
    
    for i, dataset in enumerate(datasets):
        if i + 1 < len(table.rows):
            table.rows[i + 1].cells[0].text = dataset
            for j, model in enumerate(models):
                if j + 1 < len(table.rows[i + 1].cells):
                    if dataset in main_results and model in main_results[dataset]:
                        mae = main_results[dataset][model]['MAE']
                        table.rows[i + 1].cells[j + 1].text = f"{mae:.4f}"


def update_table6(table, time_constant):
    """更新表6"""
    print("更新表6...")
    
    # 清空表格内容（保留表头）
    for i in range(1, len(table.rows)):
        for j in range(len(table.rows[i].cells)):
            table.rows[i].cells[j].text = ""
    
    # 填充数据
    for i, layer in enumerate(time_constant['layers']):
        if i + 1 < len(table.rows):
            table.rows[i + 1].cells[0].text = f"Layer {layer['layer']}"
            table.rows[i + 1].cells[1].text = f"{layer['tau_mean']:.4f}"
            table.rows[i + 1].cells[2].text = f"{layer['tau_std']:.4f}"
            table.rows[i + 1].cells[3].text = f"{layer['tau_min']:.4f}"
            table.rows[i + 1].cells[4].text = f"{layer['tau_max']:.4f}"
            table.rows[i + 1].cells[5].text = f"{layer['tau_median']:.4f}"
    
    # 填充全局统计
    if len(time_constant['layers']) + 1 < len(table.rows):
        row_idx = len(time_constant['layers']) + 1
        table.rows[row_idx].cells[0].text = "Global"
        table.rows[row_idx].cells[1].text = f"{time_constant['global']['tau_mean']:.4f}"
        table.rows[row_idx].cells[2].text = f"{time_constant['global']['tau_std']:.4f}"
        table.rows[row_idx].cells[3].text = f"{time_constant['global']['tau_min']:.4f}"
        table.rows[row_idx].cells[4].text = f"{time_constant['global']['tau_max']:.4f}"
        table.rows[row_idx].cells[5].text = f"{time_constant['global']['tau_median']:.4f}"


if __name__ == "__main__":
    update_paper_tables()
