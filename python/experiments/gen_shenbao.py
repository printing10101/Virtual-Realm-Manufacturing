"""
生成大创赛申报书
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path


def set_run_font(run, font_name='宋体', size=Pt(12), bold=False):
    """设置文本格式"""
    run.font.size = size
    run.font.bold = bold
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_heading_custom(doc, text, level=1):
    """添加自定义标题"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, '黑体', Pt(16), bold=True)
    elif level == 2:
        set_run_font(run, '黑体', Pt(14), bold=True)
    elif level == 3:
        set_run_font(run, '黑体', Pt(12), bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_body_text(doc, text, indent=True):
    """添加正文段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, '宋体', Pt(12))
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    return p


def create_application_form():
    """创建申报书"""
    doc = Document()
    
    # 设置页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    
    # ===== 封面 =====
    for _ in range(3):
        doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('大学生创新创业训练计划项目')
    set_run_font(run, '黑体', Pt(26), bold=True)
    
    doc.add_paragraph()
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('申  报  书')
    set_run_font(run, '黑体', Pt(36), bold=True)
    
    for _ in range(3):
        doc.add_paragraph()
    
    # 封面信息表
    cover_table = doc.add_table(rows=6, cols=2)
    cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover_data = [
        ('项目名称', '延迟嵌入液态神经逻辑网络用于铣削颤振稳定性预测'),
        ('项目类别', '创新训练项目'),
        ('项目负责人', '（填写）'),
        ('所在学院', '（填写）'),
        ('指导教师', '（填写）'),
        ('填报日期', '2026年06月'),
    ]
    for i, (label, value) in enumerate(cover_data):
        row = cover_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run_font(run, '宋体', Pt(14))
    
    for _ in range(4):
        doc.add_paragraph()
    
    notice = doc.add_paragraph()
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = notice.add_run('填 写 说 明')
    set_run_font(run, '黑体', Pt(14), bold=True)
    
    notice_text = """1. 申报书各项内容应认真填写，字迹清楚，表述简明。
2. 项目负责人填写基本信息和项目组其他成员。
3. 指导教师审阅并签署意见。
4. 申报书打印后由项目负责人交所在学院审核。"""
    add_body_text(doc, notice_text, indent=False)
    
    # 分页
    doc.add_page_break()
    
    # ===== 一、项目基本信息 =====
    add_heading_custom(doc, '一、项目基本信息', level=1)
    
    info_table = doc.add_table(rows=8, cols=4)
    info_table.style = 'Table Grid'
    info_data = [
        ['项目名称', '延迟嵌入液态神经逻辑网络用于铣削颤振稳定性预测', '项目英文名称', 'Delay-embedded Liquid Neural Network for Milling Chatter Stability Prediction'],
        ['所在学院', '', '联系电话', ''],
        ['项目负责人', '', '学号', ''],
        ['指导教师', '', '职称', ''],
        ['项目类别', '创新训练项目', '项目级别', '国家级'],
        ['起止时间', '2026年01月—2027年06月', '', ''],
        ['项目组成员', '（3-5人）', '', ''],
        ['申请经费', '（万元）', '', ''],
    ]
    for i, row_data in enumerate(info_data):
        row = info_table.rows[i]
        for j, cell_text in enumerate(row_data):
            row.cells[j].text = cell_text
            for p in row.cells[j].paragraphs:
                for run in p.runs:
                    set_run_font(run, '宋体', Pt(10))
    
    doc.add_paragraph()
    
    # ===== 二、项目研究背景与意义 =====
    add_heading_custom(doc, '二、项目研究背景与意义', level=1)
    
    add_heading_custom(doc, '2.1 研究背景', level=2)
    add_body_text(doc, '随着航空航天整体构件、精密模具、汽车关键零部件等高端制造领域对铣削加工质量与效率的要求日益提高，再生颤振（Regenerative Chatter）已成为制约加工精度、表面完整性与刀具寿命的核心瓶颈。颤振发生时，刀具与工件之间形成自激振动闭环，导致加工表面出现明显振纹、表面粗糙度恶化300%~500%、刀具寿命缩短60%~80%，严重时甚至引起主轴轴承损伤与加工噪声超标（>100 dB）等系统性问题。据统计，颤振问题每年给我国机械制造业造成数百亿元的经济损失。')
    
    add_body_text(doc, '避免颤振的核心手段是预测并避开不稳定的切削参数组合，即生成描述「主轴转速-极限切深」关系的稳定性叶图（Stability Lobe Diagram，SLD）。然而，长期以来铣削颤振稳定性预测面临三个根本性挑战：（1）传统解析法虽有严格的物理基础，但严重依赖精确的机床模态参数，实际车间难以直接应用；（2）现有数据驱动方法虽在数据充足时表现优异，但普遍面临「小样本过拟合、跨工况泛化差、预测违背物理规律」三大痛点；（3）既有物理引导神经网络（PINN）虽在物理一致性方面具有优势，但其物理约束仅作用于数值层，缺乏对「连续时间再生动力学」的针对性建模。')
    
    add_heading_custom(doc, '2.2 研究意义', level=2)
    add_body_text(doc, '本项目的研究意义在于解决「小样本、跨工况、可解释」三大痛点，使颤振稳定性预测在车间实际条件下可用、可信、可推广。具体而言：')
    add_body_text(doc, '（1）理论意义：首次将延迟嵌入液态时间常数网络引入铣削颤振领域，建立与再生颤振动力学严格同构的连续时间神经网络模型，为物理引导神经网络在制造领域的应用提供新的理论框架。')
    add_body_text(doc, '（2）技术意义：提出三层物理一致性损失函数与三阶段训练策略，系统性解决小样本场景下的冷启动问题与跨工况泛化难题，为高端制造场景下的物理引导AI提供可复用的方法论。')
    add_body_text(doc, '（3）应用意义：设计SLD-as-Prompt工程化诊断接口，实现「工艺员口述症状→自动反查不稳定性原因→给出参数调整建议」的端到端智能诊断，为车间实际部署提供可行路径。')
    
    # ===== 三、国内外研究现状 =====
    add_heading_custom(doc, '三、国内外研究现状', level=1)
    
    add_heading_custom(doc, '3.1 传统方法', level=2)
    add_body_text(doc, '国内学者在铣削颤振稳定性预测领域的研究大致可分为解析法、数值法和数据驱动三类。解析法基于Tlusty再生颤振理论展开，后续发展了多模态耦合、变切削深度等扩展模型，但应用前提是精确已知机床模态参数，实际车间难以准确测量。数值法包括有限差分法（FD）、半离散法（SEM）、高阶半离散法（HSDT）等，通过离散化求解时滞微分方程获得更精确的稳定性边界，但计算成本较高。')
    
    add_heading_custom(doc, '3.2 数据驱动方法', level=2)
    add_body_text(doc, '随着人工智能技术的发展，基于LSTM、Transformer等序列模型的颤振预测方法近年来受到广泛关注。但现有数据驱动方法普遍面临「小样本过拟合、跨工况泛化差、黑盒不可解释」三大共性瓶颈。')
    
    add_heading_custom(doc, '3.3 物理引导神经网络', level=2)
    add_body_text(doc, '国外学者在PINN应用方面开展了大量工作。Hasani等于2021/2022年提出液态时间常数网络（LTC），通过常微分方程求解器实现连续时间动态建模；Zhu等于2022年提出神经分段常数时滞微分方程（NPCDDE，AAAI 2022）；Raissi等于2019年开创性地将物理先验嵌入神经网络损失函数；Yu等于2022年提出梯度增强PINN（gPINN）；Guo等于2022年提出物理编码循环卷积网络（PeRCNN）。然而，现有工作普遍使用标准全连接网络或卷积网络作为主干，缺乏对颤振「连续时间再生」动力学特性的针对性建模。')
    
    add_heading_custom(doc, '3.4 现有研究不足', level=2)
    add_body_text(doc, '综合分析国内外研究现状，我们识别出三个关键不足：（1）网络结构与颤振动力学特性不匹配，现有方法无法天然捕捉颤振的连续再生特性；（2）小样本学习能力不足，缺乏针对冷启动问题的专门设计；（3）物理约束粒度不足，现有方法未能显式约束梯度层与频域层的一致性。')
    
    # ===== 四、项目研究内容与技术路线 =====
    add_heading_custom(doc, '四、项目研究内容与技术路线', level=1)
    
    add_heading_custom(doc, '4.1 研究内容', level=2)
    add_body_text(doc, '本项目提出延迟嵌入液态神经逻辑网络（DL-LNN），系统性地解决铣削颤振稳定性预测中的「小样本、跨工况、不可解释」三大痛点。主要研究内容包括：')
    add_body_text(doc, '（1）DL-LNN网络结构设计：在LTC主干网络中显式引入刀齿周期T构造延迟连接（受NPCDDE思想启发，首次用于铣削再生颤振动力学建模），结合Tlusty解析物理分支与门控融合机制，使网络天然契合「刀齿每转一圈形成再生」的连续时间机制。')
    add_body_text(doc, '（2）三层物理一致性损失函数设计：构建「数值层L_phys + 梯度层L_pcc（受gPINN思想启发）+ 频域层L_freq（本文提出）」的三层损失体系，从数值大小、曲线斜率、频谱能量三个维度联合约束预测SLD的物理一致性。')
    add_body_text(doc, '（3）三阶段训练策略与跨工况协议验证：提出「解析预训练 + 物理残差微调 + 主动学习」三阶段训练策略，在5个数据集上系统开展Leave-One-Material-Out（LOMO）与Leave-One-Condition-Out（LOCO）跨工况协议实验，并与8种主流基线方法进行对比。')
    add_body_text(doc, '（4）τ-模态参数映射律与可解释性分析：基于训练后的DL-LNN，拟合LTC时间常数τ与机床模态参数的物理映射关系，揭示网络时间常数的物理意义。')
    add_body_text(doc, '（5）SLD-as-Prompt工程化诊断接口：设计基于大语言模型的颤振智能诊断助手，实现端到端的工艺参数诊断。')
    
    add_heading_custom(doc, '4.2 技术路线', level=2)
    add_body_text(doc, '本项目技术路线如下：')
    add_body_text(doc, '第一阶段（理论建模）：基于Tlusty再生颤振理论建立解析物理模型，设计LTC主干网络与延迟嵌入机制，构建DL-LNN双分支网络架构。')
    add_body_text(doc, '第二阶段（算法开发）：实现三层物理一致性损失函数，开发三阶段训练策略，完成模型训练与超参数优化。')
    add_body_text(doc, '第三阶段（实验验证）：在5个数据集（PHM2010、NUAA、NIST、Benchmark-1、自采6061-T6）上开展主对比实验、消融实验、LOMO/LOCO跨工况实验，与8种基线方法（LSTM、GRU、Transformer、CNN、PINN、gPINN、PeRCNN、BPNN）进行系统对比。')
    add_body_text(doc, '第四阶段（工程应用）：设计SLD-as-Prompt诊断接口，开发LLM颤振智能诊断助手，进行工业案例验证。')
    
    # ===== 五、项目创新点 =====
    add_heading_custom(doc, '五、项目创新点', level=1)
    add_body_text(doc, '本项目的创新点主要体现在以下六个方面：')
    add_body_text(doc, '创新点1：首次将延迟嵌入液态时间常数网络（DL-LTC）应用于铣削颤振稳定性预测，通过在LTC神经元常微分方程中显式引入刀齿周期T构造延迟连接，实现了与再生颤振动力学机制的严格同构。')
    add_body_text(doc, '创新点2：首次提出「数值层+梯度层+频域层」三层物理一致性损失函数体系，从数值大小、曲线斜率、频谱能量三个维度联合约束预测稳定性叶图的物理一致性。')
    add_body_text(doc, '创新点3：首次提出「解析预训练+物理残差微调+主动学习」三阶段训练策略，有效解决了小样本场景下的冷启动问题。')
    add_body_text(doc, '创新点4：系统开展了LOMO和LOCO跨工况协议实验，验证了模型在未见材料和未见工况下的强泛化能力。')
    add_body_text(doc, '创新点5：首次揭示了LTC时间常数τ与机床模态参数的物理映射关系，实现了灰盒可解释性。')
    add_body_text(doc, '创新点6：创新性地设计了SLD-as-Prompt工程化诊断接口，实现了与现代大语言模型的协同应用。')
    
    # ===== 六、预期成果 =====
    add_heading_custom(doc, '六、预期成果', level=1)
    add_body_text(doc, '（1）学术论文：在Journal of Intelligent Manufacturing（IF≈5.9，JCR Q1）等国际权威期刊发表SCI论文1-2篇。')
    add_body_text(doc, '（2）软件著作权：登记「延迟嵌入液态神经逻辑网络铣削颤振稳定性预测系统」软件著作权1项。')
    add_body_text(doc, '（3）发明专利：申请「基于延迟嵌入液态时间常数网络的铣削颤振稳定性预测方法」发明专利1项。')
    add_body_text(doc, '（4）软件系统：开发一套完整的DL-LNN铣削颤振稳定性预测软件系统，支持多材料、多工况的稳定性叶图生成与智能诊断。')
    add_body_text(doc, '（5）竞赛成果：参加大学生创新创业训练计划项目成果展示，争取获得优秀项目。')
    
    # ===== 七、项目进度安排 =====
    add_heading_custom(doc, '七、项目进度安排', level=1)
    
    schedule_table = doc.add_table(rows=7, cols=3)
    schedule_table.style = 'Table Grid'
    schedule_data = [
        ['阶段', '时间', '主要工作内容'],
        ['第一阶段', '2026.01—2026.03', '文献调研、理论建模、数据集准备'],
        ['第二阶段', '2026.04—2026.06', 'DL-LNN网络设计、损失函数开发、模型训练'],
        ['第三阶段', '2026.07—2026.09', '主对比实验、消融实验、跨工况实验'],
        ['第四阶段', '2026.10—2026.12', '可解释性分析、SLD-as-Prompt接口开发'],
        ['第五阶段', '2027.01—2027.03', '论文撰写与投稿、软件著作权申请'],
        ['第六阶段', '2027.04—2027.06', '项目结题、成果展示、专利申请'],
    ]
    for i, row_data in enumerate(schedule_data):
        row = schedule_table.rows[i]
        for j, cell_text in enumerate(row_data):
            row.cells[j].text = cell_text
            for p in row.cells[j].paragraphs:
                for run in p.runs:
                    if i == 0:
                        set_run_font(run, '黑体', Pt(11), bold=True)
                    else:
                        set_run_font(run, '宋体', Pt(11))
    
    doc.add_paragraph()
    
    # ===== 八、经费预算 =====
    add_heading_custom(doc, '八、经费预算', level=1)
    
    budget_table = doc.add_table(rows=8, cols=4)
    budget_table.style = 'Table Grid'
    budget_data = [
        ['序号', '费用类别', '金额（元）', '用途说明'],
        ['1', '设备费', '5000', 'GPU服务器租赁/云计算资源'],
        ['2', '材料费', '3000', '刀具、工件材料（6061-T6、7075-T6、304SS等）'],
        ['3', '测试化验加工费', '4000', '模态测试、颤振实验数据采集'],
        ['4', '差旅费', '3000', '学术会议、竞赛展示'],
        ['5', '出版/文献/信息传播费', '2000', '论文版面费、文献数据库'],
        ['6', '其他', '3000', '软件许可、打印装订等'],
        ['', '合计', '20000', ''],
    ]
    for i, row_data in enumerate(budget_data):
        row = budget_table.rows[i]
        for j, cell_text in enumerate(row_data):
            row.cells[j].text = cell_text
            for p in row.cells[j].paragraphs:
                for run in p.runs:
                    if i == 0 or i == 7:
                        set_run_font(run, '黑体', Pt(11), bold=True)
                    else:
                        set_run_font(run, '宋体', Pt(11))
    
    doc.add_paragraph()
    
    # ===== 九、项目组成员 =====
    add_heading_custom(doc, '九、项目组成员', level=1)
    
    member_table = doc.add_table(rows=6, cols=5)
    member_table.style = 'Table Grid'
    member_data = [
        ['序号', '姓名', '学号', '所在学院/专业', '项目分工'],
        ['1', '（填写）', '', '（填写）', '项目负责人，总体设计与算法开发'],
        ['2', '（填写）', '', '（填写）', '实验数据采集与处理'],
        ['3', '（填写）', '', '（填写）', '模型训练与优化'],
        ['4', '（填写）', '', '（填写）', '软件工程化与接口开发'],
        ['5', '（填写）', '', '（填写）', '文献调研与论文撰写'],
    ]
    for i, row_data in enumerate(member_data):
        row = member_table.rows[i]
        for j, cell_text in enumerate(row_data):
            row.cells[j].text = cell_text
            for p in row.cells[j].paragraphs:
                for run in p.runs:
                    if i == 0:
                        set_run_font(run, '黑体', Pt(11), bold=True)
                    else:
                        set_run_font(run, '宋体', Pt(11))
    
    doc.add_paragraph()
    
    # ===== 十、指导教师意见 =====
    add_heading_custom(doc, '十、指导教师意见', level=1)
    
    for _ in range(6):
        doc.add_paragraph()
    
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = sig.add_run('指导教师签字：              ')
    set_run_font(run, '宋体', Pt(12))
    sig2 = doc.add_paragraph()
    sig2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = sig2.add_run('日期：    年  月  日        ')
    set_run_font(run, '宋体', Pt(12))
    
    doc.add_paragraph()
    
    # ===== 十一、学院意见 =====
    add_heading_custom(doc, '十一、学院意见', level=1)
    
    for _ in range(4):
        doc.add_paragraph()
    
    sig3 = doc.add_paragraph()
    sig3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = sig3.add_run('学院负责人签字：              ')
    set_run_font(run, '宋体', Pt(12))
    sig4 = doc.add_paragraph()
    sig4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = sig4.add_run('（学院盖章）              ')
    set_run_font(run, '宋体', Pt(12))
    sig5 = doc.add_paragraph()
    sig5.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = sig5.add_run('日期：    年  月  日        ')
    set_run_font(run, '宋体', Pt(12))
    
    # 保存文档
    output_path = Path("../../docs/大创赛/申报书.docx")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    
    print(f"✓ 申报书已生成: {output_path}")
    return output_path


if __name__ == "__main__":
    create_application_form()
