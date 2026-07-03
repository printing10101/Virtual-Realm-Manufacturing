"""
检测论文中的 ⬜ 占位符并输出待人工审阅的位置列表。

学术诚信说明（重要）：
    本脚本**不再手工填充**任何占位符内容。
    原版本通过 ``actual_data`` 字典硬编码"症状1/2/3/真实查询"等编造的工艺数据
    并自动写入论文 docx，存在学术诚信隐患（数据来源不明、数值凭空编造）。

    本重写版本仅做**检测与报告**：
      1. 扫描论文 docx 中的 ⬜ 占位符位置（段落索引、表格索引、单元格坐标）；
      2. 输出结构化 JSON 报告（占位符位置 + 上下文片段）；
      3. **不修改任何论文 docx 文件**；
      4. 占位符的实际内容应由论文作者基于真实实验结果人工填写，
         或由 ``fill_table7.py`` 派生的 Markdown/JSON 内容人工审阅后采用。

输出文件：
    - ``results/paper_placeholders_report.json`` （结构化报告）
    - ``results/paper_placeholders_report.md``  （Markdown 报告）

注意：
    原 ``paper_path`` 指向 ``docs/DL-LNN-论文-最终版-完整.docx``，
    该文件当前不存在；脚本会扫描 ``docs/`` 下所有可用的 .docx 论文文件。
"""
from __future__ import annotations

import json
import logging
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

# 项目根目录（相对本脚本位置：python/experiments/ -> 上一级 python/ -> 上一级 项目根）
_PROJECT_ROOT = Path(__file__).resolve().parents[2]
_DOCS_DIR = _PROJECT_ROOT / "docs"
_RESULTS_DIR = Path(__file__).resolve().parent / "results"
_PLACEHOLDER_MARKER = "⬜"


def _find_paper_docx_files() -> list[Path]:
    """查找 docs/ 目录下所有可用的论文 .docx 文件。"""
    if not _DOCS_DIR.exists():
        return []
    candidates: list[Path] = []
    for docx in _DOCS_DIR.rglob("*.docx"):
        # 跳过临时/备份文件
        if docx.name.startswith("~$") or docx.name.startswith("."):
            continue
        candidates.append(docx)
    return candidates


def _scan_docx_placeholders(docx_path: Path) -> dict[str, Any]:
    """扫描单个 docx 文件中的 ⬜ 占位符。

    Args:
        docx_path: docx 文件路径

    Returns:
        包含 file / paragraph_placeholders / table_placeholders /
        total_count 四键的字典
    """
    result: dict[str, Any] = {
        "file": str(docx_path.relative_to(_PROJECT_ROOT)),
        "paragraph_placeholders": [],
        "table_placeholders": [],
        "total_count": 0,
    }

    try:
        from docx import Document
    except ImportError:
        logger.warning("python-docx 未安装，无法扫描 %s", docx_path)
        result["error"] = "python-docx 未安装"
        return result

    try:
        doc = Document(str(docx_path))
    except (OSError, ValueError) as e:
        logger.warning("读取 docx 失败 (%s): %s", docx_path, e)
        result["error"] = f"读取失败: {e}"
        return result

    # 扫描段落
    for i, para in enumerate(doc.paragraphs):
        if _PLACEHOLDER_MARKER in para.text:
            result["paragraph_placeholders"].append({
                "paragraph_index": i,
                "context": para.text[:200],
            })
            result["total_count"] += 1

    # 扫描表格
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                if _PLACEHOLDER_MARKER in cell.text:
                    result["table_placeholders"].append({
                        "table_index": t_idx,
                        "row_index": r_idx,
                        "col_index": c_idx,
                        "context": cell.text[:200],
                    })
                    result["total_count"] += 1

    return result


def _render_markdown_report(
    paper_files: list[Path],
    scan_results: list[dict[str, Any]],
) -> str:
    """渲染 Markdown 报告。"""
    lines: list[str] = []
    lines.append("# 论文 ⬜ 占位符检测报告\n")
    lines.append("> **学术诚信声明**：")
    lines.append(
        "> 本脚本仅检测占位符位置，**不填充任何内容**。"
        "占位符的实际内容应由论文作者基于真实实验结果人工填写。"
    )
    lines.append("")
    lines.append(f"**扫描的论文文件数**：{len(paper_files)}\n")

    if not paper_files:
        lines.append("⚠️ 未在 `docs/` 目录下找到任何 .docx 论文文件。\n")
        return "\n".join(lines)

    total_placeholders = sum(r.get("total_count", 0) for r in scan_results)
    lines.append(f"**检测到的占位符总数**：{total_placeholders}\n")

    for scan in scan_results:
        lines.append(f"## 文件：`{scan['file']}`\n")
        if "error" in scan:
            lines.append(f"⚠️ 扫描失败：{scan['error']}\n")
            continue

        lines.append(f"- 段落占位符数：{len(scan['paragraph_placeholders'])}")
        lines.append(f"- 表格占位符数：{len(scan['table_placeholders'])}")
        lines.append(f"- 合计：{scan['total_count']}\n")

        if scan["paragraph_placeholders"]:
            lines.append("### 段落占位符位置\n")
            lines.append("| 段落索引 | 上下文（前 200 字符） |")
            lines.append("|----------|------------------------|")
            for p in scan["paragraph_placeholders"]:
                ctx = p["context"].replace("|", "\\|").replace("\n", " ")
                lines.append(
                    f"| {p['paragraph_index']} | {ctx} |"
                )
            lines.append("")

        if scan["table_placeholders"]:
            lines.append("### 表格占位符位置\n")
            lines.append("| 表格索引 | 行 | 列 | 上下文（前 200 字符） |")
            lines.append("|----------|----|----|------------------------|")
            for t in scan["table_placeholders"]:
                ctx = t["context"].replace("|", "\\|").replace("\n", " ")
                lines.append(
                    f"| {t['table_index']} | {t['row_index']} | "
                    f"{t['col_index']} | {ctx} |"
                )
            lines.append("")

    lines.append("## 处理建议\n")
    lines.append(
        "1. **不要**手工编造数据填充占位符；"
    )
    lines.append(
        "2. 对于表7（SLD-as-Prompt 模板）的占位符，"
        "参考 `fill_table7.py` 派生的 Markdown/JSON 内容；"
    )
    lines.append(
        "3. 对于其他实验数值占位符，"
        "从 `python/experiments/results/` 下的真实实验结果 JSON 中提取；"
    )
    lines.append(
        "4. 所有填充内容均需论文作者人工审阅后采用。"
    )

    return "\n".join(lines)


def main() -> None:
    """主入口：扫描论文占位符并输出报告。"""
    print("=" * 80)
    print("检测论文 ⬜ 占位符（仅报告，不填充任何内容）")
    print("=" * 80)

    paper_files = _find_paper_docx_files()
    if not paper_files:
        print("⚠️ 未在 docs/ 目录下找到任何 .docx 论文文件")

    print(f"扫描的论文文件数: {len(paper_files)}")
    for p in paper_files:
        print(f"  - {p.relative_to(_PROJECT_ROOT)}")

    scan_results = [_scan_docx_placeholders(p) for p in paper_files]
    total = sum(r.get("total_count", 0) for r in scan_results)
    print(f"\n检测到的占位符总数: {total}")

    # 输出 JSON
    _RESULTS_DIR.mkdir(parents=True, exist_ok=True)
    json_path = _RESULTS_DIR / "paper_placeholders_report.json"
    with json_path.open("w", encoding="utf-8") as f:
        json.dump(
            {
                "scanned_files": [str(p.relative_to(_PROJECT_ROOT)) for p in paper_files],
                "total_placeholders": total,
                "results": scan_results,
            },
            f,
            ensure_ascii=False,
            indent=2,
        )
    print(f"\n[OK] JSON 报告已生成: {json_path}")

    # 输出 Markdown
    md_path = _RESULTS_DIR / "paper_placeholders_report.md"
    md_content = _render_markdown_report(paper_files, scan_results)
    md_path.write_text(md_content, encoding="utf-8")
    print(f"[OK] Markdown 报告已生成: {md_path}")

    print("\n--- Markdown 报告预览 ---")
    print(md_content)
    print("=" * 80)
    print("占位符检测完成。请人工审阅报告后决定如何填写占位符。")
    print("=" * 80)


if __name__ == "__main__":
    main()
