"""
提取论文中表7的内容和位置信息
"""
from docx import Document
from pathlib import Path

paper_path = Path("../../docs/DL-LNN-论文-更新版.docx")
doc = Document(str(paper_path))

print("=" * 80)
print("搜索表7相关信息")
print("=" * 80)

# 搜索包含"表7"或"Table 7"或"SLD"的段落
for i, para in enumerate(doc.paragraphs):
    if '表7' in para.text or 'Table 7' in para.text or 'SLD-as-Prompt' in para.text or 'SLD as Prompt' in para.text:
        print(f"\n段落 {i}: {para.text[:200]}")
        # 显示前后段落
        if i > 0:
            print(f"  前一段 ({i-1}): {doc.paragraphs[i-1].text[:150]}")
        if i < len(doc.paragraphs) - 1:
            print(f"  后一段 ({i+1}): {doc.paragraphs[i+1].text[:150]}")

# 显示所有表格的简要信息
print("\n" + "=" * 80)
print("所有表格概览")
print("=" * 80)

for i, table in enumerate(doc.tables):
    table_idx = doc.element.body.index(table._element)
    
    # 向前查找最近的段落
    prev_para_text = ""
    for j in range(table_idx - 1, max(0, table_idx - 5), -1):
        if doc.element.body[j].tag.endswith('p'):
            prev_para_text = doc.paragraphs[j].text[:100]
            break
    
    print(f"\n表格 {i}: 行={len(table.rows)}, 列={len(table.columns)}")
    print(f"  前一段: {prev_para_text}")
    if len(table.rows) > 0:
        row0 = [cell.text[:15] for cell in table.rows[0].cells]
        print(f"  表头: {row0}")
