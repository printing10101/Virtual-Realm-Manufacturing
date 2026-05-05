import json
import uuid
from dataclasses import dataclass, field
from datetime import datetime
from typing import Any


@dataclass
class DocumentTemplate:
    id: str
    name: str
    description: str
    sections: list[dict[str, Any]]
    metadata: dict[str, Any] = field(default_factory=dict)


class DocumentGenerator:
    def __init__(self, llm_service=None):
        self.templates: dict[str, DocumentTemplate] = {}
        self.generated_documents: dict[str, dict[str, Any]] = {}
        self.llm_service = llm_service
        self._initialize_templates()

    def _initialize_templates(self):
        self.templates["process_card"] = DocumentTemplate(
            id="process_card",
            name="工艺卡片",
            description="A4单页工艺卡片，包含零件基本信息、工序流程表、关键工艺参数汇总",
            sections=[
                {
                    "type": "static_text",
                    "title": "零件基本信息",
                    "content": "# {{part_name}}\n\n| 项目 | 内容 |\n|------|------|\n| 物料号 | {{material_no}} |\n| 零件名称 | {{part_name}} |\n| 材料 | {{material}} |\n| 版本 | {{version}} |"
                },
                {
                    "type": "table",
                    "title": "工序流程表",
                    "data_source": "process_route",
                    "columns": [
                        {"key": "step", "header": "工序号"},
                        {"key": "operation", "header": "工序名称"},
                        {"key": "machine", "header": "设备"},
                        {"key": "description", "header": "工序说明"}
                    ],
                    "style": {"bordered": True, "striped": True, "header_color": "#409EFF"}
                },
                {
                    "type": "param_list",
                    "title": "关键工艺参数汇总",
                    "data_source": "cutting_parameters",
                    "format": "table",
                    "columns": [
                        {"key": "step", "header": "工序"},
                        {"key": "operation", "header": "操作"},
                        {"key": "v", "header": "切削速度(m/min)"},
                        {"key": "f", "header": "进给量(mm/rev)"},
                        {"key": "ap", "header": "切深(mm)"},
                        {"key": "n", "header": "转速(rpm)"}
                    ]
                },
                {
                    "type": "static_text",
                    "title": "预估工时和成本",
                    "content": "| 项目 | 数值 |\n|------|------|\n| 加工时间 | {{processing_time}} 分钟 |\n| 预估成本 | ¥{{cost}} |\n| 刀具寿命 | {{tool_life}} 分钟 |"
                },
                {
                    "type": "llm_generated",
                    "title": "备注说明",
                    "prompt": "基于以下工艺方案数据，生成一段简短的备注说明，包括工艺特点、注意事项和建议。数据：{{data}}",
                    "max_tokens": 500
                }
            ],
            metadata={"paper_size": "A4", "layout": "single_page", "template_type": "process_card"}
        )

        self.templates["work_instruction"] = DocumentTemplate(
            id="work_instruction",
            name="作业指导书",
            description="A4多页作业指导书，包含安全须知、准备工作、详细操作步骤、质量检查点和异常处理方案",
            sections=[
                {
                    "type": "static_text",
                    "title": "文档信息",
                    "content": "# {{part_name}} 作业指导书\n\n| 项目 | 内容 |\n|------|------|\n| 物料号 | {{material_no}} |\n| 零件名称 | {{part_name}} |\n| 材料 | {{material}} |\n| 版本 | {{version}} |\n| 编制日期 | {{date}} |"
                },
                {
                    "type": "llm_generated",
                    "title": "安全须知",
                    "prompt": "基于以下加工工艺（材料：{{material}}，零件类型：{{part_type}}，工序：{{process_route}}），生成详细的安全须知，包括个人防护要求、设备操作安全注意事项、紧急处理措施等。确保内容全面、无安全隐患。",
                    "max_tokens": 800
                },
                {
                    "type": "llm_generated",
                    "title": "准备工作",
                    "prompt": "基于以下工艺方案（材料：{{material}}，工序流程：{{process_route}}，设备：{{machines}}），生成加工前的准备工作清单，包括设备检查、刀具准备、材料检验、夹具准备等内容。",
                    "max_tokens": 600
                },
                {
                    "type": "llm_generated",
                    "title": "详细操作步骤",
                    "prompt": "基于以下工艺路线和参数，生成详细的编号操作步骤。工艺路线：{{process_route}}，切削参数：{{cutting_parameters}}。每个步骤需要包含：操作内容、使用的设备/刀具、关键参数、注意事项。确保步骤清晰、可执行。",
                    "max_tokens": 1500
                },
                {
                    "type": "llm_generated",
                    "title": "质量检查点",
                    "prompt": "基于以下工艺方案，识别并列出关键质量检查点。材料：{{material}}，公差要求：{{tolerance}}，表面粗糙度：{{surface_roughness}}。每个检查点需要包含：检查项目、检查时机、检测方法、合格标准。",
                    "max_tokens": 800
                },
                {
                    "type": "llm_generated",
                    "title": "常见异常处理方案",
                    "prompt": "基于以下加工工艺（材料：{{material}}，零件类型：{{part_type}}，工艺路线：{{process_route}}），列出可能的加工异常及其处理方案。包括：异常现象、可能原因、处理方法、预防措施。",
                    "max_tokens": 1000
                }
            ],
            metadata={"paper_size": "A4", "layout": "multi_page", "template_type": "work_instruction"}
        )

        self.templates["inspection_standard"] = DocumentTemplate(
            id="inspection_standard",
            name="检验标准",
            description="A4检验标准文档，包含尺寸公差表、表面质量要求、形位公差要求、检测工具清单和抽样方案",
            sections=[
                {
                    "type": "static_text",
                    "title": "文档信息",
                    "content": "# {{part_name}} 检验标准\n\n| 项目 | 内容 |\n|------|------|\n| 物料号 | {{material_no}} |\n| 零件名称 | {{part_name}} |\n| 材料 | {{material}} |\n| 版本 | {{version}} |\n| 检验标准版本 | V1.0 |"
                },
                {
                    "type": "table",
                    "title": "尺寸公差表",
                    "data_source": "dimensional_tolerances",
                    "columns": [
                        {"key": "feature", "header": "特征名称"},
                        {"key": "basic_size", "header": "基本尺寸"},
                        {"key": "tolerance_upper", "header": "上偏差"},
                        {"key": "tolerance_lower", "header": "下偏差"},
                        {"key": "unit", "header": "单位"},
                        {"key": "critical", "header": "关键项"}
                    ],
                    "style": {"bordered": True, "highlight_critical": True, "header_color": "#E6A23C"}
                },
                {
                    "type": "table",
                    "title": "表面质量要求",
                    "data_source": "surface_quality",
                    "columns": [
                        {"key": "surface", "header": "表面位置"},
                        {"key": "roughness", "header": "粗糙度要求"},
                        {"key": "grade", "header": "等级"},
                        {"key": "detection_method", "header": "检测方法"}
                    ],
                    "style": {"bordered": True, "header_color": "#67C23A"}
                },
                {
                    "type": "table",
                    "title": "形位公差要求",
                    "data_source": "geometric_tolerances",
                    "columns": [
                        {"key": "feature", "header": "特征"},
                        {"key": "tolerance_type", "header": "公差类型"},
                        {"key": "tolerance_value", "header": "公差值"},
                        {"key": "datum", "header": "基准"},
                        {"key": "unit", "header": "单位"}
                    ],
                    "style": {"bordered": True, "header_color": "#F56C6C"}
                },
                {
                    "type": "table",
                    "title": "检测工具清单",
                    "data_source": "inspection_tools",
                    "columns": [
                        {"key": "name", "header": "工具名称"},
                        {"key": "model", "header": "型号"},
                        {"key": "precision", "header": "精度"},
                        {"key": "usage", "header": "用途"}
                    ],
                    "style": {"bordered": True, "header_color": "#409EFF"}
                },
                {
                    "type": "static_text",
                    "title": "抽样方案",
                    "content": "| 项目 | 要求 |\n|------|------|\n| 抽样比例 | {{sampling_rate}} |\n| 批量范围 | {{batch_range}} |\n| 合格判定标准 | AQL {{aql_value}} |\n| 检验水平 | {{inspection_level}} |"
                }
            ],
            metadata={"paper_size": "A4", "layout": "table_based", "template_type": "inspection_standard"}
        )

    def get_template(self, template_id: str) -> DocumentTemplate | None:
        return self.templates.get(template_id)

    def list_templates(self) -> list[dict[str, Any]]:
        return [
            {
                "id": t.id,
                "name": t.name,
                "description": t.description,
                "preview_url": f"/api/v1/documents/templates/{t.id}/preview",
                "metadata": t.metadata
            }
            for t in self.templates.values()
        ]

    async def generate_document(self, template_id: str, process_plan: dict[str, Any], user_id: str) -> dict[str, Any]:
        template = self.get_template(template_id)
        if not template:
            raise ValueError(f"Template '{template_id}' not found")

        doc_id = f"doc_{uuid.uuid4().hex[:12]}"
        created_at = datetime.now().isoformat()

        rendered_content = await self._render_template(template, process_plan)

        document = {
            "doc_id": doc_id,
            "template_id": template_id,
            "template_name": template.name,
            "title": f"{process_plan.get('part_name', '未知零件')} - {template.name}",
            "content": rendered_content,
            "process_plan_id": process_plan.get("plan_id", ""),
            "user_id": user_id,
            "created_at": created_at,
            "updated_at": created_at,
            "status": "completed",
            "version": 1,
            "modifications": None
        }

        self.generated_documents[doc_id] = document

        return {
            "doc_id": doc_id,
            "status": "completed",
            "estimated_time": 2
        }

    async def _render_template(self, template: DocumentTemplate, process_plan: dict[str, Any]) -> str:
        rendered_sections = []

        for section in template.sections:
            rendered = await self._render_section(section, process_plan)
            if rendered:
                rendered_sections.append(rendered)

        return "\n\n".join(rendered_sections)

    async def _render_section(self, section: dict[str, Any], process_plan: dict[str, Any]) -> str | None:
        section_type = section.get("type")

        if section_type == "static_text":
            return self._render_static_text(section, process_plan)
        elif section_type == "table":
            return self._render_table(section, process_plan)
        elif section_type == "param_list":
            return self._render_param_list(section, process_plan)
        elif section_type == "llm_generated":
            return await self._render_llm_generated(section, process_plan)

        return None

    def _render_static_text(self, section: dict[str, Any], process_plan: dict[str, Any]) -> str:
        content = section.get("content", "")
        content = self._fill_placeholders(content, process_plan)

        title = section.get("title", "")
        return f"## {title}\n\n{content}"

    def _render_table(self, section: dict[str, Any], process_plan: dict[str, Any]) -> str:
        title = section.get("title", "")
        data_source = section.get("data_source", "")
        columns = section.get("columns", [])

        table_data = process_plan.get(data_source, [])
        if not table_data:
            return f"## {title}\n\n*无数据*"

        header = "| " + " | ".join(col["header"] for col in columns) + " |"
        separator = "| " + " | ".join("---" for _ in columns) + " |"

        rows = []
        for item in table_data:
            if isinstance(item, dict):
                row = "| " + " | ".join(str(item.get(col["key"], "")) for col in columns) + " |"
                rows.append(row)
            elif isinstance(item, list):
                row = "| " + " | ".join(str(val) for val in item[:len(columns)]) + " |"
                rows.append(row)

        table_md = f"\n{header}\n{separator}\n" + "\n".join(rows) + "\n"

        return f"## {title}{table_md}"

    def _render_param_list(self, section: dict[str, Any], process_plan: dict[str, Any]) -> str:
        title = section.get("title", "")
        data_source = section.get("data_source", "")
        columns = section.get("columns", [])

        param_data = process_plan.get(data_source, [])
        if isinstance(param_data, dict):
            param_data = param_data.get("parameters", [])

        if not param_data:
            return f"## {title}\n\n*无参数数据*"

        header = "| " + " | ".join(col["header"] for col in columns) + " |"
        separator = "| " + " | ".join("---" for _ in columns) + " |"

        rows = []
        for item in param_data:
            if isinstance(item, dict):
                row = "| " + " | ".join(str(item.get(col["key"], "")) for col in columns) + " |"
                rows.append(row)

        table_md = f"\n{header}\n{separator}\n" + "\n".join(rows) + "\n"

        return f"## {title}{table_md}"

    async def _render_llm_generated(self, section: dict[str, Any], process_plan: dict[str, Any]) -> str | None:
        title = section.get("title", "")
        prompt_template = section.get("prompt", "")

        prompt = self._fill_placeholders(prompt_template, process_plan)

        if self.llm_service:
            try:
                llm_content = await self._call_llm(prompt, section.get("max_tokens", 500))
                return f"## {title}\n\n{llm_content}"
            except Exception:
                return self._generate_fallback_content(title, process_plan)
        else:
            return self._generate_fallback_content(title, process_plan)

    async def _call_llm(self, prompt: str, max_tokens: int) -> str:
        if hasattr(self.llm_service, 'generate'):
            response = await self.llm_service.generate(
                prompt=prompt,
                max_tokens=max_tokens,
                temperature=0.7
            )
            if isinstance(response, dict):
                return response.get("content", response.get("text", ""))
            return str(response)
        elif hasattr(self.llm_service, 'chat'):
            response = await self.llm_service.chat(
                messages=[{"role": "user", "content": prompt}],
                max_tokens=max_tokens
            )
            if isinstance(response, dict):
                return response.get("content", response.get("text", ""))
            return str(response)

        return ""

    def _generate_fallback_content(self, title: str, process_plan: dict[str, Any]) -> str:
        fallbacks = {
            "备注说明": f"本工艺方案基于{process_plan.get('material', '指定材料')}材料设计，适用于{process_plan.get('part_type', '该类型')}零件加工。建议加工过程中严格控制切削参数，确保加工质量。",
            "安全须知": f"加工{process_plan.get('material', '该材料')}材料时，请遵守以下安全规定：\n1. 佩戴防护眼镜和手套\n2. 确保设备接地良好\n3. 加工前检查刀具状态\n4. 遵守设备操作规程",
            "准备工作": f"加工前准备工作：\n1. 检查{process_plan.get('material', '材料')}材料规格\n2. 准备所需刀具和夹具\n3. 校准加工设备\n4. 确认工艺参数设置",
            "详细操作步骤": f"根据工艺路线执行以下操作：\n{self._format_process_route_fallback(process_plan.get('process_route', []))}",
            "质量检查点": f"质量检查要求：\n1. 尺寸公差检查：{process_plan.get('tolerance', '按图纸要求')}\n2. 表面粗糙度检查：{process_plan.get('surface_roughness', '按工艺要求')}\n3. 形位公差检查：按图纸标注",
            "常见异常处理方案": "常见异常处理：\n1. 表面粗糙度超标：检查刀具磨损，调整切削参数\n2. 尺寸超差：检查夹具定位，补偿刀具磨损\n3. 刀具异常磨损：降低切削速度，检查冷却液"
        }

        content = fallbacks.get(title, f"{title}内容待补充。")
        return content

    def _format_process_route_fallback(self, process_route: list) -> str:
        if not process_route:
            return "按工艺卡片工序顺序执行。"
        lines = []
        for step in process_route:
            if isinstance(step, dict):
                lines.append(f"{step.get('step', '?')}. {step.get('operation', '?')}: {step.get('description', '?')}")
        return "\n".join(lines) if lines else "按工艺卡片工序顺序执行。"

    def _fill_placeholders(self, text: str, process_plan: dict[str, Any]) -> str:
        import re

        def replace_placeholder(match):
            key = match.group(1)
            value = process_plan.get(key, f"{{{key}}}")
            if isinstance(value, list):
                return json.dumps(value, ensure_ascii=False)
            return str(value)

        return re.sub(r'\{\{(\w+)\}\}', replace_placeholder, text)

    def get_document(self, doc_id: str) -> dict[str, Any] | None:
        return self.generated_documents.get(doc_id)

    def update_document(self, doc_id: str, new_content: str, user_id: str) -> dict[str, Any] | None:
        document = self.generated_documents.get(doc_id)
        if not document:
            return None

        document["content"] = new_content
        document["updated_at"] = datetime.now().isoformat()
        document["modifications"] = {
            "modified_by": user_id,
            "modified_at": document["updated_at"],
            "original_content": document.get("content", "")
        }
        document["version"] = document.get("version", 1) + 1

        return document

    def get_document_history(self, process_plan_id: str) -> list[dict[str, Any]]:
        return [
            doc for doc in self.generated_documents.values()
            if doc.get("process_plan_id") == process_plan_id
        ]

    def export_to_pdf_data(self, doc_id: str) -> bytes | None:
        document = self.generated_documents.get(doc_id)
        if not document:
            return None

        import io
        import re

        from reportlab.lib.colors import HexColor
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm, mm
        from reportlab.platypus import (
            HRFlowable,
            Paragraph,
            SimpleDocTemplate,
            Spacer,
        )

        buffer = io.BytesIO()

        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2 * cm,
            leftMargin=2 * cm,
            topMargin=2 * cm,
            bottomMargin=2 * cm,
            title=document.get("title", "工艺文档")
        )

        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=styles['Title'],
            fontSize=18,
            textColor=HexColor('#1a1a1a'),
            spaceAfter=10 * mm,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        ))
        styles.add(ParagraphStyle(
            name='CustomHeading1',
            parent=styles['Heading1'],
            fontSize=16,
            textColor=HexColor('#409EFF'),
            spaceBefore=15 * mm,
            spaceAfter=8 * mm,
            borderWidth=1,
            borderColor=HexColor('#409EFF'),
            borderPadding=4,
            fontName='Helvetica-Bold'
        ))
        styles.add(ParagraphStyle(
            name='CustomHeading2',
            parent=styles['Heading2'],
            fontSize=13,
            textColor=HexColor('#303133'),
            spaceBefore=10 * mm,
            spaceAfter=6 * mm,
            fontName='Helvetica-Bold'
        ))
        styles.add(ParagraphStyle(
            name='CustomBody',
            parent=styles['Normal'],
            fontSize=10,
            leading=14,
            spaceAfter=6 * mm,
            alignment=TA_JUSTIFY,
            fontName='Helvetica'
        ))

        story = []

        story.append(Paragraph(document.get("title", ""), styles['CustomTitle']))
        story.append(HRFlowable(width="100%", thickness=1, color=HexColor('#409EFF'), spaceAfter=5 * mm))

        story.append(Paragraph(f"<b>模板：</b>{document.get('template_name', '')}", styles['CustomBody']))
        story.append(Paragraph(f"<b>生成时间：</b>{document.get('created_at', '')[:19]}", styles['CustomBody']))
        story.append(Spacer(1, 5 * mm))

        content = document.get("content", "")

        sections = re.split(r'\n## ', content)

        for i, section in enumerate(sections):
            if i == 0 and not section.startswith('#'):
                continue

            if section.startswith('# '):
                section = section[2:]
            elif section.startswith('## '):
                section = section[3:]

            lines = section.strip().split('\n', 1)
            if len(lines) == 1:
                story.append(Paragraph(lines[0], styles['CustomHeading1']))
                story.append(Spacer(1, 3 * mm))
                continue

            title, body = lines
            story.append(Paragraph(title, styles['CustomHeading2']))

            if '| ' in body and '\n' in body:
                table_element = self._parse_markdown_table(body)
                if table_element:
                    story.append(table_element)
                    story.append(Spacer(1, 4 * mm))
                else:
                    story.append(Paragraph(body.replace('\n', '<br/>'), styles['CustomBody']))
            else:
                formatted_body = body.replace('\n', '<br/>')
                story.append(Paragraph(formatted_body, styles['CustomBody']))

            story.append(Spacer(1, 2 * mm))

        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()

    def _parse_markdown_table(self, table_text: str) -> Any | None:
        from reportlab.lib.colors import HexColor, white
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Table, TableStyle

        lines = [line.strip() for line in table_text.strip().split('\n') if line.strip()]
        if len(lines) < 2:
            return None

        data = []
        for line in lines:
            if '|---' in line:
                continue
            cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            if cells:
                data.append(cells)

        if not data:
            return None

        col_count = max(len(row) for row in data)
        for row in data:
            while len(row) < col_count:
                row.append('')

        styles = getSampleStyleSheet()
        from reportlab.lib.styles import ParagraphStyle
        cell_style = ParagraphStyle('TableCell', parent=styles['Normal'], fontSize=9, leading=12)

        styled_data = []
        for row in data:
            styled_row = [Paragraph(cell, cell_style) for cell in row]
            styled_data.append(styled_row)

        col_widths = [70] * col_count
        table = Table(styled_data, colWidths=col_widths)

        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), HexColor('#409EFF')),
            ('TEXTCOLOR', (0, 0), (-1, 0), white),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
            ('BACKGROUND', (0, 1), (-1, -1), white),
            ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#dcdfe6')),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [white, HexColor('#f5f7fa')]),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        ]))

        return table

    def export_to_docx_data(self, doc_id: str) -> bytes | None:
        document = self.generated_documents.get(doc_id)
        if not document:
            return None

        import io
        import re

        from docx import Document as DocxDocument
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt, RGBColor

        doc = DocxDocument()

        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        title = doc.add_heading(document.get("title", ""), level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run1 = meta.add_run(f"模板：{document.get('template_name', '')}    ")
        run1.font.size = Pt(10)
        run1.font.color.rgb = RGBColor(0x60, 0x62, 0x66)
        run2 = meta.add_run(f"生成时间：{document.get('created_at', '')[:19]}")
        run2.font.size = Pt(10)
        run2.font.color.rgb = RGBColor(0x60, 0x62, 0x66)

        content = document.get("content", "")
        sections = re.split(r'\n## ', content)

        for i, section in enumerate(sections):
            if i == 0 and not section.startswith('#'):
                continue

            if section.startswith('# '):
                section = section[2:]
            elif section.startswith('## '):
                section = section[3:]

            lines = section.strip().split('\n', 1)
            if len(lines) == 1:
                heading = doc.add_heading(lines[0], level=1)
                for run in heading.runs:
                    run.font.color.rgb = RGBColor(0x40, 0x9E, 0xFF)
                continue

            title_text, body = lines

            heading = doc.add_heading(title_text, level=2)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0x30, 0x31, 0x33)

            if '| ' in body and '\n' in body:
                self._add_markdown_table_to_docx(doc, body)
            else:
                paragraphs = body.split('\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        p = doc.add_paragraph()
                        run = p.add_run(para_text)
                        run.font.size = Pt(11)
                        p.paragraph_format.space_after = Pt(6)

            doc.add_paragraph()

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _add_markdown_table_to_docx(self, doc, table_text: str):
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor

        lines = [line.strip() for line in table_text.strip().split('\n') if line.strip()]
        if len(lines) < 2:
            return

        data = []
        for line in lines:
            if '|---' in line:
                continue
            cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            if cells:
                data.append(cells)

        if not data:
            return

        col_count = max(len(row) for row in data)
        for row in data:
            while len(row) < col_count:
                row.append('')

        table = doc.add_table(rows=len(data), cols=col_count)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Light Shading Accent 1'

        for i, row_data in enumerate(data):
            for j, cell_text in enumerate(row_data):
                cell = table.cell(i, j)
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    paragraph.alignment = 1
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        if i == 0:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for row in table.rows:
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = qn('w:trHeight')
            from lxml import etree
            height_elem = etree.SubElement(trPr, trHeight)
            height_elem.set(qn('w:val'), "350")

    def duplicate_document(self, doc_id: str, user_id: str) -> dict[str, Any] | None:
        original = self.generated_documents.get(doc_id)
        if not original:
            return None

        new_doc_id = f"doc_{uuid.uuid4().hex[:12]}"
        new_doc = original.copy()
        new_doc["doc_id"] = new_doc_id
        new_doc["created_at"] = datetime.now().isoformat()
        new_doc["updated_at"] = new_doc["created_at"]
        new_doc["version"] = 1
        new_doc["user_id"] = user_id
        new_doc["source_doc_id"] = doc_id
        new_doc["modifications"] = None

        self.generated_documents[new_doc_id] = new_doc

        return new_doc
